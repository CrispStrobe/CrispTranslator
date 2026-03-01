#!/usr/bin/env python3
"""
Format Transplant
=================
Apply the complete layout/formatting of a blueprint DOCX (1) to the content
of a source DOCX (2). Produces a new document that has (2)'s text formatted
exactly according to (1)'s layout system.

Strategy (Approach B):
  Start from a copy of blueprint (1), clear its body, then re-insert content
  from (2) with (1)'s styles applied.  Page layout, style definitions,
  headers, and footers all come from (1).  Text content and inline formatting
  (bold/italic/underline) come from (2).

Usage:
  python format_transplant.py blueprint.docx source.docx output.docx [-v]
  python format_transplant.py blueprint.docx source.docx output.docx \\
      --style-map "My Body=Normal" "My Heading=Heading 1"
"""

import argparse
import asyncio
import copy
import logging
import os
import re
import shutil
import sys
import time
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple

# ============================================================================
# SYSTEM CHECK
# ============================================================================

print("Format Transplant – System Check")
print("-" * 44)


def _check(name: str, stmt: str) -> bool:
    try:
        exec(stmt, globals())
        print(f"  ✓ {name}")
        return True
    except ImportError as e:
        print(f"  ✗ {name}: {e}")
        return False
    except Exception as e:
        print(f"  ✗ {name} (unexpected): {e}")
        return False


HAS_DOCX = _check(
    "python-docx",
    "from docx import Document; from docx.shared import Pt, RGBColor, Emu; "
    "from docx.text.paragraph import Paragraph; "
    "from docx.oxml.shared import OxmlElement; from docx.oxml.ns import qn; "
    "from docx.oxml import parse_xml",
)
HAS_LXML  = _check("lxml",       "from lxml import etree")
HAS_OPENAI    = _check("openai",       "from openai import OpenAI")
HAS_ANTHROPIC = _check("anthropic",    "import anthropic")
HAS_POE       = _check("fastapi-poe",  "import fastapi_poe as fp")
HAS_REQUESTS  = _check("requests",     "import requests")

print("-" * 44)

if not HAS_DOCX or not HAS_LXML:
    print("ERROR: Required libraries missing.  Install with:")
    print("  pip install python-docx lxml")
    sys.exit(1)

# Safe imports after checks
from docx import Document  # noqa: E402
from docx.oxml import parse_xml  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.oxml.shared import OxmlElement  # noqa: E402
from docx.shared import Emu, Pt, RGBColor  # noqa: E402
from docx.text.paragraph import Paragraph  # noqa: E402
from lxml import etree  # noqa: E402
import requests  # noqa: E402

# ============================================================================
# LOGGING
# ============================================================================

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)-7s] %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger("FormatTransplant")

def load_dotenv(path: Optional[Path] = None):
    """Simple .env loader to avoid extra dependencies."""
    env_path = path or Path(".env")
    if not env_path.exists():
        return
    try:
        with open(env_path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("#"):
                    continue
                if "=" in line:
                    key, _, value = line.partition("=")
                    # Strip quotes if present
                    value = value.strip().strip('"').strip("'")
                    os.environ[key.strip()] = value
    except Exception as e:
        logger.warning(f"Failed to load .env: {e}")

# Load environment early
load_dotenv()

# ============================================================================
# SEMANTIC CLASSIFICATION CONSTANTS
# ============================================================================

# Run-level XML tags to KEEP (semantic inline formatting) when cleaning rPr.
# Everything else (fonts, sizes, colors, language, kern) gets stripped so
# the blueprint style governs the visual appearance.
KEEP_RPR_TAGS: Set[str] = {
    qn("w:b"),
    qn("w:bCs"),
    qn("w:i"),
    qn("w:iCs"),
    qn("w:u"),
    qn("w:strike"),
    qn("w:dstrike"),
    qn("w:vertAlign"),
    qn("w:highlight"),
    qn("w:smallCaps"),
    qn("w:allCaps"),
    qn("w:em"),
    qn("w:vanish"),
}

# Multilingual heading style name patterns (lowercase), grouped by level.
HEADING_PATTERNS: Dict[int, List[str]] = {
    1: [
        "heading 1", "heading1", "h1",
        "überschrift 1", "titre 1", "titolo 1", "encabezado 1",
        "заголовок 1", "标题 1", "kop 1", "nagłówek 1", "rubrik 1",
        "heading1char",
    ],
    2: [
        "heading 2", "heading2", "h2",
        "überschrift 2", "titre 2", "titolo 2", "encabezado 2",
        "заголовок 2", "标题 2", "kop 2", "nagłówek 2",
    ],
    3: [
        "heading 3", "heading3", "h3",
        "überschrift 3", "titre 3", "titolo 3", "encabezado 3",
        "заголовок 3", "标题 3", "kop 3", "nagłówek 3",
    ],
    4: ["heading 4", "heading4", "h4", "überschrift 4", "titre 4", "заголовок 4"],
    5: ["heading 5", "heading5", "h5", "überschrift 5", "titre 5"],
    6: ["heading 6", "heading6", "h6", "überschrift 6"],
    7: ["heading 7", "heading7", "h7"],
    8: ["heading 8", "heading8", "h8"],
    9: ["heading 9", "heading9", "h9"],
}

TITLE_PATTERNS = ["title", "documenttitle", "thetitle", "doc title"]
BODY_PATTERNS = [
    "normal", "standard", "body text", "bodytext", "fließtext",
    "texte de corps", "corpo del testo", "cuerpo de texto",
    "основной текст", "no spacing", "default paragraph style", "tekst podstawowy",
]
FOOTNOTE_PATTERNS = [
    "footnote text", "fußnotentext", "note de bas de page",
    "nota a piè di pagina", "nota al pie", "сноска",
    "footnote", "footnotetext",
]
CAPTION_PATTERNS = ["caption", "bildunterschrift", "légende", "didascalia", "leyenda"]
BLOCKQUOTE_PATTERNS = [
    "block text", "blockquote", "quote", "intense quote",
    "block quotation", "zitat", "citation", "citazione", "bloque de texto",
]
ABSTRACT_PATTERNS = ["abstract", "zusammenfassung", "résumé", "riassunto"]

# Regex to detect "Heading_02", "Ueberschrift_01", "Titre2", "Titolo3" etc.
# Matches: <heading-keyword> <optional separators> <1-digit level>
# Used as fallback when exact/prefix HEADING_PATTERNS don't match.
_HEADING_KW_RE = re.compile(
    r"(?:heading|ueberschrift|\u00fcberschrift|titre|titolo|encabezado"
    r"|\u0437\u0430\u0433\u043e\u043b\u043e\u0432\u043e\u043a"  # заголовок
    r"|kop|rubrik|nag\u0142\u00f3wek"                           # nagłówek
    r")[\s_\-]*0*([1-9])",
    re.IGNORECASE | re.UNICODE,
)

# ============================================================================
# DATA STRUCTURES
# ============================================================================


@dataclass
class RunData:
    """Data captured from a single DOCX run."""

    text: str
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[bool] = None
    font_name: Optional[str] = None
    font_size_pt: Optional[float] = None
    font_color: Optional[Tuple[int, int, int]] = None
    # Raw lxml element – needed for footnote-reference runs and for deep copy
    raw_xml: Optional[Any] = None


@dataclass
class ParagraphData:
    """Data captured from a single DOCX paragraph."""

    runs: List[RunData] = field(default_factory=list)
    original_style_name: str = "Normal"
    semantic_class: str = "body"   # body / heading1..9 / title / footnote / caption / blockquote / table
    heading_level: int = 0         # 1-9 for headings, 0 otherwise

    # Direct paragraph formatting from source (informational; blueprint style overrides)
    alignment: Optional[Any] = None
    left_indent_pt: Optional[float] = None
    right_indent_pt: Optional[float] = None
    first_line_indent_pt: Optional[float] = None
    space_before_pt: Optional[float] = None
    space_after_pt: Optional[float] = None
    line_spacing: Optional[Any] = None

    location: str = "body"         # body / table / footnote / header / footer
    footnote_id: Optional[str] = None

    # Raw lxml <w:p> element (for deep-copy strategy)
    raw_xml: Optional[Any] = None

    # Structural flags
    has_footnote_refs: bool = False
    is_section_break: bool = False   # inline <w:sectPr> inside <w:pPr>

    def get_text(self) -> str:
        return "".join(r.text for r in self.runs)


@dataclass
class FootnoteData:
    """Data from a single footnote."""

    footnote_id: str
    paragraphs: List[ParagraphData] = field(default_factory=list)
    raw_xml: Optional[Any] = None   # The <w:footnote> element


@dataclass
class BlueprintStyleInfo:
    """Resolved information about one style in the blueprint."""

    name: str
    style_id: str
    type_val: int   # 1=paragraph, 2=character, 3=table, 4=numbering
    base_style_name: Optional[str] = None
    resolved_font: Optional[str] = None
    resolved_size_pt: Optional[float] = None
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    left_indent_pt: Optional[float] = None
    space_before_pt: Optional[float] = None
    space_after_pt: Optional[float] = None
    # OOXML outline level (0=H1 … 8=H9); None if not a heading style
    outline_level: Optional[int] = None


@dataclass
class BlueprintSchema:
    """Full formatting schema extracted from the blueprint document."""

    sections: List[Dict[str, Any]] = field(default_factory=list)
    styles: Dict[str, BlueprintStyleInfo] = field(default_factory=dict)       # name → info
    style_id_map: Dict[str, BlueprintStyleInfo] = field(default_factory=dict) # id → info
    default_font: str = "Times New Roman"
    default_font_size_pt: float = 12.0
    # Style names that actually appear in the blueprint body (for diagnostics)
    body_para_style_names: Set[str] = field(default_factory=set)
    # Character style ID used for footnote number runs (e.g. "FootnoteReference")
    footnote_ref_char_style_id: str = "FootnoteReference"
    # Actual <w:rPr> element deep-copied from the blueprint's own footnote marker
    # runs. Applied verbatim so font, size, and superscript match the blueprint.
    # None = blueprint has no numbered footnotes (fall back to char style reference).
    footnote_marker_rPr_xml: Optional[Any] = None
    # Separator text that the blueprint places immediately after the footnote number
    # (typically "\t", sometimes " ", rarely ""). None = not yet determined.
    footnote_separator: Optional[str] = None


# ============================================================================
# LLM CONFIGURATION
# ============================================================================

class LLMProvider(Enum):
    OPENAI     = "openai"
    ANTHROPIC  = "anthropic"
    GROQ       = "groq"
    NEBIUS     = "nebius"
    SCALEWAY   = "scaleway"
    OPENROUTER = "openrouter"
    MISTRAL    = "mistral"
    POE        = "poe"
    OLLAMA     = "ollama"


# Per-provider defaults — base_url=None means the provider uses its own SDK
# Added top 5 fallback models for each provider
PROVIDER_DEFAULTS: Dict[str, Dict[str, Any]] = {
    "openai": {
        "base_url": "https://api.openai.com/v1",
        "env": "OPENAI_API_KEY",
        "model": "gpt-4o",
        "fallbacks": ["gpt-4o-2024-08-06", "gpt-4o-mini", "gpt-4-turbo", "gpt-3.5-turbo"],
        "batch_size": 15
    },
    "anthropic": {
        "base_url": None,
        "env": "ANTHROPIC_API_KEY",
        "model": "claude-3-7-sonnet-20250219",
        "fallbacks": ["claude-3-5-sonnet-20241022", "claude-3-5-haiku-20241022", "claude-3-opus-20240229", "claude-2.1"],
        "batch_size": 15
    },
    "groq": {
        "base_url": "https://api.groq.com/openai/v1",
        "env": "GROQ_API_KEY",
        "model": "llama-3.3-70b-versatile",
        "fallbacks": ["llama-3.1-70b-versatile", "llama-3.1-8b-instant", "mixtral-8x7b-32768", "gemma2-9b-it"],
        "batch_size": 5
    },
    "nebius": {
        "base_url": "https://api.studio.nebius.ai/v1",
        "env": "NEBIUS_API_KEY",
        "model": "meta-llama/Meta-Llama-3.1-70B-Instruct",
        "fallbacks": ["meta-llama/Meta-Llama-3.1-8B-Instruct", "meta-llama/Llama-Guard-3-8B"],
        "batch_size": 15
    },
    "scaleway": {
        "base_url": "https://api.scaleway.ai/v1",
        "env": "SCALEWAY_API_KEY", # Updated to match .env
        "model": "llama-3.3-70b-instruct",
        "fallbacks": ["deepseek-r1-distill-llama-70b", "llama-3.1-8b-instruct", "mistral-nemo-instruct-2407", "pixtral-12b-2409"],
        "batch_size": 15
    },
    "openrouter": {
        "base_url": "https://openrouter.ai/api/v1",
        "env": "OPENROUTER_API_KEY",
        "model": "meta-llama/llama-3.3-70b-instruct",
        "fallbacks": ["anthropic/claude-3.5-sonnet", "google/gemini-pro-1.5", "mistralai/mistral-large", "qwen/qwen-2.5-72b-instruct"],
        "batch_size": 15
    },
    "mistral": {
        "base_url": "https://api.mistral.ai/v1",
        "env": "MISTRAL_API_KEY",
        "model": "mistral-large-latest",
        "fallbacks": ["mistral-medium-latest", "mistral-small-latest", "codestral-latest", "open-mistral-nemo"],
        "batch_size": 15
    },
    "poe": {
        "base_url": None,
        "env": "POE_API_KEY",
        "model": "Claude-3.7-Sonnet",
        "fallbacks": ["Claude-3.5-Sonnet", "GPT-4o", "Claude-3-Opus", "Llama-3.1-405B"],
        "batch_size": 15
    },
    "ollama": {
        "base_url": "http://localhost:11434/api",
        "env": "OLLAMA_API_KEY",
        "model": "llama3.2",
        "fallbacks": ["llama3.1", "mistral", "phi3", "qwen2.5"],
        "batch_size": 15
    },
}


@dataclass
class LLMConfig:
    """Runtime configuration for an LLM provider."""
    provider: LLMProvider
    model: str
    api_key: str
    base_url: Optional[str] = None     # overrides PROVIDER_DEFAULTS if set
    max_tokens: int = 4096
    temperature: float = 0.1           # low for deterministic formatting
    # How many chars of blueprint text to send for styleguide generation (~10 K tokens)
    blueprint_context_chars: int = 100_000
    # Source paragraphs per LLM batch
    para_batch_size: int = 15
    # Retry settings
    max_retries: int = 5
    retry_delay_s: float = 5.0
    fallback_models: List[str] = field(default_factory=list)


def llm_config_from_args(
    provider_str: str,
    model: Optional[str] = None,
    api_key: Optional[str] = None,
) -> LLMConfig:
    """Build an LLMConfig from CLI/UI inputs, filling defaults from PROVIDER_DEFAULTS."""
    import os
    defaults = PROVIDER_DEFAULTS.get(provider_str, {})
    resolved_key = api_key or os.getenv(defaults.get("env", ""), "")
    if not resolved_key:
        raise ValueError(
            f"No API key for provider '{provider_str}'. "
            f"Set env var {defaults.get('env', '?')} or pass --llm-key."
        )
    
    # Handle 'auto' or 'default' markers from UI/CLI
    resolved_model = model
    if resolved_model and resolved_model.lower() in ("auto", "default"):
        resolved_model = None
        
    return LLMConfig(
        provider=LLMProvider(provider_str),
        model=resolved_model or defaults.get("model", ""),
        api_key=resolved_key,
        base_url=defaults.get("base_url"),
        para_batch_size=defaults.get("batch_size", 15),
        fallback_models=defaults.get("fallbacks", []),
    )


# ============================================================================
# UTILITY HELPERS
# ============================================================================

_W_NS  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"
_W_NS_MAP = {"w": _W_NS}

# Paragraph-level attributes that encode source-document revision session IDs
# (rsidR, etc.) and Word 2010+ paragraph identity GUIDs (w14:paraId/textId).
# These come from the *source* document but settings.xml comes from the
# *blueprint*, so the rsid values are absent from <w:rsids> in settings.xml.
# Word treats that mismatch as "unreadable content". Strip them all.
_PARA_STRIP_ATTRS: Set[str] = {
    f"{{{_W14_NS}}}paraId",
    f"{{{_W14_NS}}}textId",
    f"{{{_W_NS}}}rsidR",
    f"{{{_W_NS}}}rsidRPr",
    f"{{{_W_NS}}}rsidDel",
    f"{{{_W_NS}}}rsidRDefault",
    f"{{{_W_NS}}}rsidRPrChange",
}
_RUN_STRIP_ATTRS: Set[str] = {
    f"{{{_W_NS}}}rsidRPr",
    f"{{{_W_NS}}}rsidDel",
}


def _w(tag: str) -> str:
    """Return Clark-notation tag for namespace 'w'."""
    return f"{{{_W_NS}}}{tag}"


def _xpath(elem: Any, xpath_str: str) -> list:
    """
    XPath helper that works on both python-docx BaseOxmlElement instances
    (which override .xpath() to accept no kwargs) and plain lxml _Element
    instances (e.g. from parse_xml on footnotes XML).
    Bypasses the python-docx override by calling the underlying lxml method
    directly, always supplying the 'w:' namespace binding.
    """
    return etree._Element.xpath(elem, xpath_str, namespaces=_W_NS_MAP)


def _strip_tracking_attrs(elem: Any) -> None:
    """
    Strip source-document revision-tracking attributes from every <w:p> and
    <w:r> node in the subtree (including the root element itself).

    Root cause of "Word found unreadable content": paragraphs deep-copied from
    the source carry rsidR/w14:paraId values that reference revision sessions
    recorded in the source's settings.xml. The output document's settings.xml
    comes from the blueprint, so those session IDs are absent from <w:rsids>.
    Word flags the mismatch. Stripping the attributes entirely is safe —
    rsid tracking is optional and Word regenerates them on next save.
    """
    p_tag = f"{{{_W_NS}}}p"
    r_tag = f"{{{_W_NS}}}r"
    for node in elem.iter():
        if node.tag == p_tag:
            for attr in _PARA_STRIP_ATTRS:
                node.attrib.pop(attr, None)
        elif node.tag == r_tag:
            for attr in _RUN_STRIP_ATTRS:
                node.attrib.pop(attr, None)


def classify_style(style_name: str) -> Tuple[str, int]:
    """
    Classify a style name into (semantic_class, heading_level).

    Returns e.g. ("heading3", 3) or ("body", 0) or ("footnote", 0).
    """
    name_lo = style_name.lower().strip()

    # Headings — exact / prefix match
    for level, patterns in HEADING_PATTERNS.items():
        if name_lo in patterns:
            return f"heading{level}", level
        # Prefix match handles "Heading 1 Char" etc.
        for pat in patterns:
            if name_lo.startswith(pat):
                return f"heading{level}", level

    # Headings — regex: catches "Ueberschrift_01", "Titre2", "MyHeading3" …
    m = _HEADING_KW_RE.search(name_lo)
    if m:
        level = int(m.group(1))
        return f"heading{level}", level

    # Title
    if name_lo in TITLE_PATTERNS:
        return "title", 0

    # Other semantic classes
    for pat in FOOTNOTE_PATTERNS:
        if pat in name_lo:
            return "footnote", 0
    for pat in CAPTION_PATTERNS:
        if pat in name_lo:
            return "caption", 0
    for pat in BLOCKQUOTE_PATTERNS:
        if pat in name_lo:
            return "blockquote", 0
    for pat in ABSTRACT_PATTERNS:
        if pat in name_lo:
            return "abstract", 0
    for pat in BODY_PATTERNS:
        if name_lo == pat or name_lo.startswith(pat):
            return "body", 0

    return "unknown", 0


def _has_footnote_ref(p_elem: Any) -> bool:
    return bool(_xpath(p_elem, ".//w:footnoteReference | .//w:footnoteRef"))


def _has_inline_sect_pr(p_elem: Any) -> bool:
    return bool(_xpath(p_elem, "./w:pPr/w:sectPr"))


# ============================================================================
# PHASE 1 – BLUEPRINT ANALYZER
# ============================================================================


class BlueprintAnalyzer:
    """
    Deeply analyses the blueprint document (1) to extract its complete
    formatting schema: page layout, all style definitions, body para inventory.
    """

    def analyze(self, doc: Document) -> BlueprintSchema:
        logger.info("[BLUEPRINT] ══════════════════════════════════════════")
        logger.info("[BLUEPRINT] Analysing blueprint document…")
        schema = BlueprintSchema()
        self._sections(doc, schema)
        self._styles(doc, schema)
        self._defaults(doc, schema)
        self._body_inventory(doc, schema)
        self._footnote_format(doc, schema)
        logger.info(
            "[BLUEPRINT] Done: %d section(s), %d style(s), "
            "%d unique body-para styles",
            len(schema.sections),
            len(schema.styles),
            len(schema.body_para_style_names),
        )
        return schema

    # ------------------------------------------------------------------
    def _sections(self, doc: Document, schema: BlueprintSchema) -> None:
        logger.debug("[BLUEPRINT] ── Sections ──")
        for i, sect in enumerate(doc.sections):
            try:
                def _pt(v):
                    return round(v.pt, 2) if v is not None else None

                data = {
                    "index": i,
                    "page_width_pt":      _pt(sect.page_width),
                    "page_height_pt":     _pt(sect.page_height),
                    "left_margin_pt":     _pt(sect.left_margin),
                    "right_margin_pt":    _pt(sect.right_margin),
                    "top_margin_pt":      _pt(sect.top_margin),
                    "bottom_margin_pt":   _pt(sect.bottom_margin),
                    "gutter_pt":          _pt(sect.gutter),
                    "header_distance_pt": _pt(sect.header_distance),
                    "footer_distance_pt": _pt(sect.footer_distance),
                    "orientation":        str(sect.orientation),
                    "diff_first_page":    sect.different_first_page_header_footer,
                }
                schema.sections.append(data)
                logger.debug(
                    "[BLUEPRINT] Section %d | %.0fx%.0f pt | "
                    "Margins L:%.0f R:%.0f T:%.0f B:%.0f | "
                    "Header-dist:%.0f Footer-dist:%.0f | Gutter:%.0f",
                    i,
                    data["page_width_pt"] or 0,
                    data["page_height_pt"] or 0,
                    data["left_margin_pt"] or 0,
                    data["right_margin_pt"] or 0,
                    data["top_margin_pt"] or 0,
                    data["bottom_margin_pt"] or 0,
                    data["header_distance_pt"] or 0,
                    data["footer_distance_pt"] or 0,
                    data["gutter_pt"] or 0,
                )
            except Exception as exc:
                logger.warning("[BLUEPRINT] Section %d error: %s", i, exc)

    # ------------------------------------------------------------------
    def _styles(self, doc: Document, schema: BlueprintSchema) -> None:
        logger.debug("[BLUEPRINT] ── Styles ──")
        _fn_ref_style_found = False   # track whether we've already locked in the style
        for style in doc.styles:
            try:
                # Skip numbering styles — they have no font/base_style attributes
                if not hasattr(style, "font"):
                    continue
                type_val = style.type.value if hasattr(style.type, "value") else int(style.type)
                info = BlueprintStyleInfo(
                    name=style.name,
                    style_id=style.style_id,
                    type_val=type_val,
                )
                base = getattr(style, "base_style", None)
                if base:
                    info.base_style_name = base.name

                info.resolved_font = self._resolve_font(style)
                info.resolved_size_pt = self._resolve_size(style)

                if style.font:
                    info.bold = style.font.bold
                    info.italic = style.font.italic

                # Extract outline level from raw XML (language-independent)
                try:
                    ol_el = style._element.find(f".//{qn('w:outlineLvl')}")
                    if ol_el is not None:
                        ol_val = ol_el.get(qn("w:val"), "")
                        if ol_val.isdigit():
                            info.outline_level = int(ol_val)
                except Exception:
                    pass

                # Detect blueprint's footnote-reference character style (first match wins)
                if type_val == 2 and not _fn_ref_style_found:
                    nm_key = (
                        style.name.lower()
                        .replace(" ", "").replace("-", "").replace("_", "")
                    )
                    if any(kw in nm_key for kw in (
                        "footnotereference", "funotenzeichen",
                        "fußnotenzeichen", "fu\u00dfnotenzeichen",
                        "noteref", "notefnref",
                    )):
                        schema.footnote_ref_char_style_id = style.style_id
                        _fn_ref_style_found = True
                        logger.debug(
                            "[BLUEPRINT] Footnote reference char style: '%s' → id='%s'",
                            style.name, style.style_id,
                        )

                if type_val == 1 and hasattr(style, "paragraph_format"):
                    pf = style.paragraph_format
                    try:
                        info.left_indent_pt = pf.left_indent.pt if pf.left_indent else None
                        info.space_before_pt = pf.space_before.pt if pf.space_before else None
                        info.space_after_pt = pf.space_after.pt if pf.space_after else None
                    except Exception:
                        pass

                schema.styles[style.name] = info
                schema.style_id_map[style.style_id] = info

                if type_val == 1:
                    logger.debug(
                        "[BLUEPRINT] ParaStyle %-32s  id=%-20s  font=%-18s  "
                        "%.0fpt  bold=%-5s  italic=%-5s  base=%s",
                        f"'{style.name}'",
                        f"'{style.style_id}'",
                        f"'{info.resolved_font}'",
                        info.resolved_size_pt or 0,
                        info.bold,
                        info.italic,
                        f"'{info.base_style_name}'",
                    )
            except Exception as exc:
                logger.warning(
                    "[BLUEPRINT] Style '%s' error: %s",
                    getattr(style, "name", "?"),
                    exc,
                )

    # ------------------------------------------------------------------
    def _resolve_font(self, style) -> Optional[str]:
        """Walk up style hierarchy to find the first explicitly set font name."""
        curr = style
        while curr is not None:
            try:
                if curr.font and curr.font.name:
                    return curr.font.name
            except Exception:
                pass
            curr = getattr(curr, "base_style", None)
        return None

    def _resolve_size(self, style) -> Optional[float]:
        """Walk up style hierarchy to find the first explicitly set font size."""
        curr = style
        while curr is not None:
            try:
                if curr.font and curr.font.size:
                    return curr.font.size.pt
            except Exception:
                pass
            curr = getattr(curr, "base_style", None)
        return None

    # ------------------------------------------------------------------
    def _defaults(self, doc: Document, schema: BlueprintSchema) -> None:
        try:
            normal = doc.styles["Normal"]
            schema.default_font = self._resolve_font(normal) or "Times New Roman"
            schema.default_font_size_pt = self._resolve_size(normal) or 12.0
        except Exception as exc:
            logger.warning("[BLUEPRINT] Could not resolve default font: %s", exc)
        logger.debug(
            "[BLUEPRINT] Document defaults: font='%s'  size=%.0fpt",
            schema.default_font,
            schema.default_font_size_pt,
        )

    # ------------------------------------------------------------------
    def _body_inventory(self, doc: Document, schema: BlueprintSchema) -> None:
        for para in doc.paragraphs:
            try:
                schema.body_para_style_names.add(para.style.name)
            except Exception:
                pass
        logger.debug(
            "[BLUEPRINT] Body para styles present: %s",
            sorted(schema.body_para_style_names),
        )

    # ------------------------------------------------------------------
    def _footnote_format(self, doc: Document, schema: BlueprintSchema) -> None:
        """
        Read the first 3 blueprint footnotes to learn the exact formatting the
        blueprint uses for footnote marker runs and the separator that follows them.

        Two things are extracted:
          footnote_marker_rPr_xml — the <w:rPr> element from the <w:footnoteRef>
              run, deep-copied verbatim. Captures font name, size, vertAlign,
              superscript, color etc. exactly as they appear in the blueprint.
          footnote_separator — the text content of the run immediately after the
              marker run: "\t" (tab), " " (space), "" (none), or anything else.

        Both are read from the *blueprint's own footnotes* (not the source),
        so the output always matches the blueprint's convention regardless of
        what the source document was doing.
        """
        try:
            fn_part = None
            for rel in doc.part.rels.values():
                if "relationships/footnotes" in rel.reltype:
                    fn_part = rel.target_part
                    break
            if fn_part is None:
                logger.debug(
                    "[BLUEPRINT] No footnotes part – footnote format detection skipped"
                )
                return

            root = parse_xml(fn_part.blob)
            rPr_found = False
            sep_found = False
            samples = 0

            for fn_elem in _xpath(root, "//w:footnote"):
                try:
                    fn_id = int(fn_elem.get(_w("id"), "0"))
                except (ValueError, TypeError):
                    continue
                if fn_id <= 0:
                    continue  # Word-internal separators / continuation markers

                samples += 1
                if samples > 3:
                    break

                # Only the first paragraph of each footnote carries the marker
                p_elems = _xpath(fn_elem, ".//w:p")
                if not p_elems:
                    continue
                p_elem = p_elems[0]
                runs = list(p_elem.findall(qn("w:r")))

                for ri, r_elem in enumerate(runs):
                    if not _xpath(r_elem, ".//w:footnoteRef"):
                        continue

                    # ── Marker rPr (verbatim deep-copy) ──────────────────
                    if not rPr_found:
                        rPr = r_elem.find(qn("w:rPr"))
                        if rPr is not None:
                            schema.footnote_marker_rPr_xml = copy.deepcopy(rPr)
                            rPr_found = True
                            logger.debug(
                                "[BLUEPRINT] Footnote marker rPr captured "
                                "(fn id=%d): %s",
                                fn_id,
                                [c.tag.split("}")[-1] for c in rPr],
                            )
                        else:
                            logger.debug(
                                "[BLUEPRINT] Footnote marker run has no rPr (fn id=%d)",
                                fn_id,
                            )

                    # ── Separator after marker ────────────────────────────
                    # A separator run is one whose ENTIRE text content is
                    # whitespace (tab, space, or empty) OR contains a <w:tab/>.
                    # If the next run has actual content, this footnote has no
                    # dedicated separator run — skip it and try the next footnote.
                    if not sep_found:
                        if ri + 1 < len(runs):
                            next_r = runs[ri + 1]
                            has_tab = next_r.find(qn("w:tab")) is not None
                            t_elems = next_r.findall(qn("w:t"))
                            sep_text = "".join(t.text or "" for t in t_elems)
                            
                            if has_tab:
                                # Prioritize physical tab element over text
                                schema.footnote_separator = "\t"
                                sep_found = True
                                logger.debug(
                                    "[BLUEPRINT] Footnote separator: <w:tab/> (fn id=%d)",
                                    fn_id,
                                )
                            elif sep_text.strip() == "":
                                # Pure whitespace → this IS the separator run
                                schema.footnote_separator = sep_text
                                sep_found = True
                                label = repr(sep_text) if sep_text else "(empty)"
                                logger.debug(
                                    "[BLUEPRINT] Footnote separator: %s (fn id=%d)",
                                    label, fn_id,
                                )
                            else:
                                # Next run is actual footnote text — no separator
                                # run in this footnote; keep looking in later ones
                                logger.debug(
                                    "[BLUEPRINT] Footnote id=%d: no separator run "
                                    "(text starts immediately after marker)",
                                    fn_id,
                                )
                        # else: no run after marker — keep looking

                    break  # found the marker in this footnote; move to next footnote

                if rPr_found and sep_found:
                    break

            if samples == 0:
                logger.debug("[BLUEPRINT] Blueprint has no numbered footnotes to sample")
            else:
                # If we sampled footnotes but never found a pure-whitespace separator
                # run, the blueprint uses no separator — record that explicitly.
                if not sep_found:
                    schema.footnote_separator = ""
                    logger.debug(
                        "[BLUEPRINT] No separator run found across %d sampled footnote(s)"
                        " — blueprint uses no explicit separator",
                        samples,
                    )
                logger.info(
                    "[BLUEPRINT] Footnote format: marker_rPr=%s  separator=%s",
                    "captured" if rPr_found else "none",
                    repr(schema.footnote_separator)
                    if schema.footnote_separator is not None
                    else "not found",
                )

        except Exception as exc:
            logger.warning(
                "[BLUEPRINT] Footnote format detection error: %s", exc, exc_info=True
            )


# ============================================================================
# PHASE 2 – CONTENT EXTRACTOR
# ============================================================================


class ContentExtractor:
    """
    Extracts all content from the source document (2), preserving text and
    inline semantic formatting (bold/italic/underline).  Direct paragraph
    formatting values are recorded for debug purposes but are NOT applied
    to the output – the blueprint style governs layout.
    """

    def __init__(self) -> None:
        # Built during extraction: source style_id → style_name
        self.src_style_id_to_name: Dict[str, str] = {}

    # ------------------------------------------------------------------
    def extract(
        self, doc: Document
    ) -> Tuple[List[ParagraphData], List[FootnoteData]]:
        """
        Returns:
            body_elements  – ordered list of ParagraphData (paragraphs AND
                             table placeholders with semantic_class='table')
            footnotes      – list of FootnoteData
        """
        logger.info("[EXTRACT] ══════════════════════════════════════════")
        logger.info("[EXTRACT] Extracting content from source document…")

        # Build source style ID→name lookup
        for s in doc.styles:
            try:
                self.src_style_id_to_name[s.style_id] = s.name
            except Exception:
                pass
        logger.debug(
            "[EXTRACT] Source document has %d styles", len(self.src_style_id_to_name)
        )

        body_elements = self._body(doc)
        footnotes = self._footnotes(doc)

        para_count = sum(1 for e in body_elements if e.semantic_class != "table")
        table_count = sum(1 for e in body_elements if e.semantic_class == "table")
        logger.info(
            "[EXTRACT] Done: %d paragraphs, %d tables, %d footnotes",
            para_count,
            table_count,
            len(footnotes),
        )
        return body_elements, footnotes

    # ------------------------------------------------------------------
    def _run(self, run) -> RunData:
        rd = RunData(text=run.text, raw_xml=run._element)
        rd.bold = run.bold
        rd.italic = run.italic
        rd.underline = run.underline
        try:
            if run.font.name:
                rd.font_name = run.font.name
            if run.font.size:
                rd.font_size_pt = run.font.size.pt
            if run.font.color and run.font.color.type is not None:
                try:
                    rgb = run.font.color.rgb
                    rd.font_color = (rgb[0], rgb[1], rgb[2])
                except Exception:
                    pass
        except Exception:
            pass
        return rd

    def _para(self, para: Paragraph, location: str = "body") -> ParagraphData:
        pd = ParagraphData(location=location, raw_xml=para._element)

        try:
            pd.original_style_name = para.style.name if para.style else "Normal"
        except Exception:
            pd.original_style_name = "Normal"

        pd.semantic_class, pd.heading_level = classify_style(pd.original_style_name)

        try:
            pd.alignment = para.alignment
            pf = para.paragraph_format
            pd.left_indent_pt = pf.left_indent.pt if pf.left_indent else None
            pd.right_indent_pt = pf.right_indent.pt if pf.right_indent else None
            pd.first_line_indent_pt = (
                pf.first_line_indent.pt if pf.first_line_indent else None
            )
            pd.space_before_pt = pf.space_before.pt if pf.space_before else None
            pd.space_after_pt = pf.space_after.pt if pf.space_after else None
            pd.line_spacing = pf.line_spacing
        except Exception as exc:
            logger.debug("[EXTRACT] Para format read error: %s", exc)

        for run in para.runs:
            try:
                pd.runs.append(self._run(run))
            except Exception as exc:
                logger.debug("[EXTRACT] Run error: %s", exc)

        pd.has_footnote_refs = _has_footnote_ref(para._element)
        pd.is_section_break = _has_inline_sect_pr(para._element)

        logger.debug(
            "[EXTRACT] Para | style='%s'  class=%s  loc=%s | "
            "runs=%d  fnRef=%s  sectBrk=%s | text='%.60s'",
            pd.original_style_name,
            pd.semantic_class,
            location,
            len(pd.runs),
            pd.has_footnote_refs,
            pd.is_section_break,
            pd.get_text(),
        )
        return pd

    # ------------------------------------------------------------------
    def _body(self, doc: Document) -> List[ParagraphData]:
        elements: List[ParagraphData] = []
        body = doc.element.body

        for child in body:
            tag = child.tag
            if tag == qn("w:p"):
                try:
                    para = Paragraph(child, doc)
                    elements.append(self._para(para, "body"))
                except Exception as exc:
                    logger.warning("[EXTRACT] Body para error: %s", exc)
            elif tag == qn("w:tbl"):
                # Table placeholder – raw XML carried along for deep copy
                placeholder = ParagraphData(
                    location="table_placeholder",
                    raw_xml=child,
                    original_style_name="__TABLE__",
                    semantic_class="table",
                )
                elements.append(placeholder)
                logger.debug("[EXTRACT] Table placeholder recorded")
            elif tag == qn("w:sectPr"):
                logger.debug(
                    "[EXTRACT] Body <w:sectPr> found (document-level) – blueprint's will be used"
                )

        self._infer_headings(elements)
        return elements

    # ------------------------------------------------------------------
    def _infer_headings(self, elements: List[ParagraphData]) -> None:
        """
        Post-process extracted body paragraphs to infer heading hierarchy
        from direct paragraph formatting when no explicit heading style exists.

        Signals:
          • All text runs bold  OR  pPr/rPr contains <w:b> (paragraph default bold)
          • Short text (< 100 chars) — headings are rarely long sentences
          • Font size: larger sizes → higher priority (lower heading level number)

        Font sizes of heading candidates are clustered descending so that:
          largest size → heading level 1
          next size    → heading level 2
          etc.
        If all candidates share the same (or no) font size, all become level 1.
        Paragraphs already classified as a non-body class are skipped.
        """
        from collections import Counter

        candidates: List[Tuple[ParagraphData, float]] = []   # (pd, size_pt)
        body_sizes: List[float] = []

        for pd in elements:
            if pd.semantic_class != "body":
                continue
            p_elem = pd.raw_xml
            text = pd.get_text().strip()
            if not text:
                continue

            # ── Paragraph-default bold and font size (from pPr/rPr) ──────────
            ppr_bold   = False
            ppr_sz_pt: Optional[float] = None
            if p_elem is not None:
                pPr = p_elem.find(qn("w:pPr"))
                if pPr is not None:
                    ppr_rPr = pPr.find(qn("w:rPr"))
                    if ppr_rPr is not None:
                        ppr_bold = ppr_rPr.find(qn("w:b")) is not None
                        sz_el = ppr_rPr.find(qn("w:sz"))
                        if sz_el is not None:
                            try:
                                ppr_sz_pt = int(sz_el.get(qn("w:val"), "0")) / 2.0
                            except (ValueError, TypeError):
                                pass

            # ── Run-level bold and font size ──────────────────────────────────
            text_runs = [rd for rd in pd.runs if rd.text.strip()]
            all_runs_bold = bool(text_runs) and all(
                rd.bold is True or (rd.bold is None and ppr_bold)
                for rd in text_runs
            )
            effective_bold = all_runs_bold or ppr_bold

            run_szs = [rd.font_size_pt for rd in text_runs if rd.font_size_pt]
            effective_sz = (sum(run_szs) / len(run_szs)) if run_szs else ppr_sz_pt

            # ── Classify ──────────────────────────────────────────────────────
            if effective_bold and 0 < len(text) < 100:
                candidates.append((pd, effective_sz or 0.0))
            elif effective_sz:
                body_sizes.append(effective_sz)

        if not candidates:
            return

        # Body text reference size (mode)
        body_sz = Counter(body_sizes).most_common(1)[0][0] if body_sizes else 0.0

        # Unique heading candidate sizes, largest first
        unique_szs = sorted({sz for _, sz in candidates if sz > 0}, reverse=True)
        # Drop sizes that are ≤ body size (same-size bold = not really a heading)
        heading_szs = [sz for sz in unique_szs if body_sz == 0.0 or sz > body_sz + 0.4]
        if not heading_szs:
            heading_szs = [0.0]   # sentinel: all candidates → level 1

        def _level(sz: float) -> int:
            if heading_szs == [0.0]:
                return 1
            for lvl, threshold in enumerate(heading_szs, start=1):
                if sz >= threshold - 0.4:
                    return lvl
            return len(heading_szs)

        for pd, sz in candidates:
            lvl = _level(sz)
            pd.semantic_class = f"heading{lvl}"
            pd.heading_level = lvl
            logger.debug(
                "[EXTRACT] Inferred heading%d (sz=%.1fpt) | '%.60s'",
                lvl, sz, pd.get_text(),
            )

        size_labels = (
            {i + 1: f"{sz:.1f}pt" for i, sz in enumerate(heading_szs)}
            if heading_szs != [0.0] else {1: "any"}
        )
        logger.info(
            "[EXTRACT] Heading inference: %d candidate(s) across %d level(s): %s",
            len(candidates), len(heading_szs), size_labels,
        )

    # ------------------------------------------------------------------
    def _footnotes(self, doc: Document) -> List[FootnoteData]:
        result: List[FootnoteData] = []
        try:
            fn_part = None
            for rel in doc.part.rels.values():
                if "relationships/footnotes" in rel.reltype:
                    fn_part = rel.target_part
                    break
            if fn_part is None:
                logger.debug("[EXTRACT] No footnotes part in source document")
                return result

            root = parse_xml(fn_part.blob)

            for fn_elem in _xpath(root, "//w:footnote"):
                fn_id = fn_elem.get(_w("id"), "0")
                if int(fn_id) <= 0:
                    continue  # Word-internal continuation separators

                fd = FootnoteData(footnote_id=fn_id, raw_xml=fn_elem)
                for p_elem in _xpath(fn_elem, ".//w:p"):
                    try:
                        para = Paragraph(p_elem, doc)
                        pd = self._para(para, "footnote")
                        pd.footnote_id = fn_id
                        fd.paragraphs.append(pd)
                    except Exception as exc:
                        logger.debug("[EXTRACT] Footnote para error: %s", exc)

                result.append(fd)
                preview = fd.paragraphs[0].get_text()[:60] if fd.paragraphs else ""
                logger.debug(
                    "[EXTRACT] Footnote id=%s: %d para(s) | '%.60s'",
                    fn_id,
                    len(fd.paragraphs),
                    preview,
                )
        except Exception as exc:
            logger.warning("[EXTRACT] Footnote extraction error: %s", exc)

        return result


# ============================================================================
# PHASE 3 – STYLE MAPPER
# ============================================================================


class StyleMapper:
    """
    Maps source document style names to blueprint style names.

    Resolution order:
      1. User-supplied explicit overrides (--style-map)
      2. Exact name match in blueprint
      3. Case-insensitive name match
      4. Semantic class match (heading level, body, footnote, caption, …)
      5. Fallback to blueprint 'Normal' (or first available para style)
    """

    def __init__(
        self,
        schema: BlueprintSchema,
        user_overrides: Optional[Dict[str, str]] = None,
    ) -> None:
        self.schema = schema
        self.user_overrides: Dict[str, str] = user_overrides or {}
        self._cache: Dict[str, str] = {}

        # Semantic lookup tables (built from blueprint)
        self._bp_headings: Dict[int, str] = {}
        self._bp_title: Optional[str] = None
        self._bp_body: Optional[str] = None
        self._bp_footnote: Optional[str] = None
        self._bp_caption: Optional[str] = None
        self._bp_blockquote: Optional[str] = None
        self._bp_abstract: Optional[str] = None

        self._build_lookup()
        self._log_lookup()

    # ------------------------------------------------------------------
    def _build_lookup(self) -> None:
        # Pass 1 — outline level from style XML (most reliable, language-independent)
        # Prefer styles actually used in the blueprint body when there are ties.
        for name, info in self.schema.styles.items():
            if info.type_val != 1 or info.outline_level is None:
                continue
            level = info.outline_level + 1   # OOXML is 0-based; TOCHeading uses 9
            if level < 1 or level > 9:
                continue
            used_first = name in self.schema.body_para_style_names
            if level not in self._bp_headings or used_first:
                self._bp_headings[level] = name
                logger.debug(
                    "[MAPPER] Blueprint heading%d from outlineLvl: '%s'", level, name
                )

        # Pass 2 — semantic name classification (fills gaps & non-heading classes)
        for name, info in self.schema.styles.items():
            if info.type_val != 1:
                continue
            sem, level = classify_style(name)

            if sem == "title" and not self._bp_title:
                self._bp_title = name
            elif sem.startswith("heading") and level > 0:
                if level not in self._bp_headings:
                    self._bp_headings[level] = name
                elif name in self.schema.body_para_style_names:
                    # Prefer actually-used style
                    self._bp_headings[level] = name
            elif sem == "body" and not self._bp_body:
                self._bp_body = name
            elif sem == "footnote" and not self._bp_footnote:
                self._bp_footnote = name
            elif sem == "caption" and not self._bp_caption:
                self._bp_caption = name
            elif sem == "blockquote" and not self._bp_blockquote:
                self._bp_blockquote = name
            elif sem == "abstract" and not self._bp_abstract:
                self._bp_abstract = name

        # Fallback body style
        if not self._bp_body:
            if "Normal" in self.schema.styles:
                self._bp_body = "Normal"
            else:
                para_styles = [
                    n for n, i in self.schema.styles.items() if i.type_val == 1
                ]
                self._bp_body = para_styles[0] if para_styles else "Normal"

    def _log_lookup(self) -> None:
        logger.info("[MAPPER] ══════════════════════════════════════════")
        logger.info("[MAPPER] Blueprint semantic lookup:")
        logger.info("[MAPPER]   Title       → '%s'", self._bp_title)
        logger.info("[MAPPER]   Headings    → %s", self._bp_headings)
        logger.info("[MAPPER]   Body        → '%s'", self._bp_body)
        logger.info("[MAPPER]   Footnote    → '%s'", self._bp_footnote)
        logger.info("[MAPPER]   Caption     → '%s'", self._bp_caption)
        logger.info("[MAPPER]   Blockquote  → '%s'", self._bp_blockquote)
        logger.info("[MAPPER]   Abstract    → '%s'", self._bp_abstract)

    # ------------------------------------------------------------------
    def map(self, src_name: str, sem_class: str, heading_level: int) -> str:
        key = f"{src_name}::{sem_class}::{heading_level}"
        if key in self._cache:
            return self._cache[key]
        result = self._resolve(src_name, sem_class, heading_level)
        self._cache[key] = result
        logger.debug(
            "[MAPPER] '%s' (class=%s hl=%d) → '%s'",
            src_name,
            sem_class,
            heading_level,
            result,
        )
        return result

    def _resolve(self, src_name: str, sem_class: str, heading_level: int) -> str:
        bp = self.schema.styles

        # 1. User override (always wins)
        if src_name in self.user_overrides:
            target = self.user_overrides[src_name]
            if target in bp:
                logger.debug("[MAPPER] User override: '%s' → '%s'", src_name, target)
                return target
            logger.warning(
                "[MAPPER] Override target '%s' not found in blueprint – ignoring",
                target,
            )

        # 2a. Semantic heading match — runs BEFORE name lookup so that paragraphs
        #     reclassified by _infer_headings (e.g. "Normal" paragraphs detected
        #     as bold+short) get the blueprint heading style, not "Normal".
        if sem_class.startswith("heading") and heading_level > 0:
            if heading_level in self._bp_headings:
                return self._bp_headings[heading_level]
            for delta in [1, -1, 2, -2, 3, -3]:
                adj = heading_level + delta
                if adj in self._bp_headings:
                    logger.debug(
                        "[MAPPER] Heading %d not in blueprint, using adjacent level %d",
                        heading_level, adj,
                    )
                    return self._bp_headings[adj]
            if self._bp_headings:
                return next(iter(self._bp_headings.values()))

        # 2b. Exact name match (for non-heading classes)
        if src_name in bp:
            return src_name

        # 3. Case-insensitive name match
        src_lo = src_name.lower()
        for bp_name in bp:
            if bp_name.lower() == src_lo:
                logger.debug(
                    "[MAPPER] Case-insensitive match: '%s' → '%s'", src_name, bp_name
                )
                return bp_name

        # 4. Semantic class match (non-heading classes; headings already handled above)
        if sem_class.startswith("heading") and heading_level > 0:
            if heading_level in self._bp_headings:
                return self._bp_headings[heading_level]
            # Try adjacent heading levels
            for delta in [1, -1, 2, -2, 3, -3]:
                adj = heading_level + delta
                if adj in self._bp_headings:
                    logger.debug(
                        "[MAPPER] Heading %d not in blueprint, using adjacent level %d",
                        heading_level,
                        adj,
                    )
                    return self._bp_headings[adj]
            if self._bp_headings:
                return next(iter(self._bp_headings.values()))

        if sem_class == "title":
            if self._bp_title:
                return self._bp_title
            if 1 in self._bp_headings:
                return self._bp_headings[1]

        if sem_class == "footnote" and self._bp_footnote:
            return self._bp_footnote

        if sem_class == "caption" and self._bp_caption:
            return self._bp_caption

        if sem_class == "blockquote" and self._bp_blockquote:
            return self._bp_blockquote

        if sem_class == "abstract" and self._bp_abstract:
            return self._bp_abstract

        # 5. Fallback
        logger.debug(
            "[MAPPER] No match for '%s' (class=%s) – falling back to '%s'",
            src_name,
            sem_class,
            self._bp_body,
        )
        return self._bp_body or "Normal"

    # ------------------------------------------------------------------
    def log_full_table(self, elements: List[ParagraphData]) -> None:
        logger.info("[MAPPER] ── Full style mapping table ──")
        seen: Dict[str, str] = {}
        for pd in elements:
            if pd.semantic_class == "table":
                continue
            key = pd.original_style_name
            if key not in seen:
                mapped = self.map(pd.original_style_name, pd.semantic_class, pd.heading_level)
                seen[key] = mapped
                logger.info("[MAPPER]   %-40s → '%s'", f"'{key}'", mapped)
        logger.info("[MAPPER] %d unique source style(s) mapped", len(seen))


# ============================================================================
# PHASE 4 – DOCUMENT BUILDER
# ============================================================================


class DocumentBuilder:
    """
    Assembles the output document:
      1. Copy blueprint (1) → output (preserves styles.xml, page layout, etc.)
      2. Clear body content (keep final <w:sectPr>)
      3. Re-insert source elements with blueprint styles
      4. Transplant footnote content
    """

    def __init__(self, schema: BlueprintSchema, mapper: StyleMapper) -> None:
        self.schema = schema
        self.mapper = mapper
        # Source style-id→name lookup, injected after extraction
        self.src_style_id_to_name: Dict[str, str] = {}
        # Optional LLM-formatted text maps: id(ParagraphData) → markdown string
        self.llm_para_map: Dict[int, str] = {}
        self.llm_fn_map:   Dict[int, str] = {}

    # ------------------------------------------------------------------
    def build(
        self,
        blueprint_path: Path,
        output_path: Path,
        body_elements: List[ParagraphData],
        footnotes: List[FootnoteData],
    ) -> None:
        logger.info("[BUILD] ══════════════════════════════════════════")
        logger.info("[BUILD] Building output document…")

        # ── Step 1: copy blueprint ──────────────────────────────────────
        shutil.copy2(str(blueprint_path), str(output_path))
        logger.info("[BUILD] Copied blueprint → %s", output_path)

        # ── Step 2: open copy ──────────────────────────────────────────
        doc = Document(str(output_path))
        self._log_doc_info(doc, "BLUEPRINT COPY (before clear)")

        # ── Step 3: clear body ─────────────────────────────────────────
        self._clear_body(doc)

        # ── Step 4: insert source elements ────────────────────────────
        self._insert_elements(doc, body_elements)

        # ── Step 5: transplant footnotes ──────────────────────────────
        if footnotes:
            self._transplant_footnotes(doc, footnotes)
        else:
            logger.info("[BUILD] No footnotes to transplant")

        # ── Step 6: save ──────────────────────────────────────────────
        doc.save(str(output_path))
        logger.info("[BUILD] ✓ Saved → %s", output_path)

        # ── Step 7: verification log ─────────────────────────────────
        try:
            verify_doc = Document(str(output_path))
            self._log_doc_info(verify_doc, "OUTPUT (verification)")
        except Exception as exc:
            logger.warning("[BUILD] Verification read failed: %s", exc)

    # ------------------------------------------------------------------
    def _log_doc_info(self, doc: Document, label: str) -> None:
        logger.debug("[BUILD] ── Doc info [%s] ──", label)
        for i, sect in enumerate(doc.sections):
            try:
                logger.debug(
                    "[BUILD]   Section %d: %.0fx%.0f pt | margins L:%.0f R:%.0f T:%.0f B:%.0f",
                    i,
                    sect.page_width.pt if sect.page_width else 0,
                    sect.page_height.pt if sect.page_height else 0,
                    sect.left_margin.pt if sect.left_margin else 0,
                    sect.right_margin.pt if sect.right_margin else 0,
                    sect.top_margin.pt if sect.top_margin else 0,
                    sect.bottom_margin.pt if sect.bottom_margin else 0,
                )
            except Exception:
                pass
        style_names = sorted(
            {p.style.name for p in doc.paragraphs if p.style}
        )
        logger.debug("[BUILD]   Body para styles present: %s", style_names)

    # ------------------------------------------------------------------
    def _clear_body(self, doc: Document) -> None:
        body = doc.element.body

        # Locate the final <w:sectPr> (document-level section properties)
        final_sect_pr = None
        for child in reversed(list(body)):
            if child.tag == qn("w:sectPr"):
                final_sect_pr = child
                break

        n_before = len(list(body))
        for child in list(body):
            if child is not final_sect_pr:
                body.remove(child)
        n_after = len(list(body))

        logger.debug(
            "[BUILD] Body cleared: %d → %d element(s) (sectPr preserved=%s)",
            n_before,
            n_after,
            final_sect_pr is not None,
        )

    # ------------------------------------------------------------------
    def _insert_elements(self, doc: Document, elements: List[ParagraphData]) -> None:
        body = doc.element.body

        # Insertion point: just before the final <w:sectPr>
        children = list(body)
        final_sect_pr = next(
            (c for c in reversed(children) if c.tag == qn("w:sectPr")), None
        )
        insert_at = children.index(final_sect_pr) if final_sect_pr is not None else len(children)

        inserted = 0
        skipped = 0

        for idx, elem in enumerate(elements):
            try:
                if elem.semantic_class == "table":
                    tbl_xml = copy.deepcopy(elem.raw_xml)
                    _strip_tracking_attrs(tbl_xml)
                    self._remap_table_styles(tbl_xml, doc)
                    body.insert(insert_at + inserted, tbl_xml)
                    inserted += 1
                    logger.debug("[BUILD] [%d] Table inserted", idx)
                else:
                    p_elem = self._build_para(elem, doc)
                    if p_elem is not None:
                        body.insert(insert_at + inserted, p_elem)
                        inserted += 1
                        target_style = self.mapper.map(
                            elem.original_style_name, elem.semantic_class, elem.heading_level
                        )
                        logger.debug(
                            "[BUILD] [%d] Para | '%s' → '%s' | '%.55s'",
                            idx,
                            elem.original_style_name,
                            target_style,
                            elem.get_text(),
                        )
                    else:
                        skipped += 1
            except Exception as exc:
                logger.error("[BUILD] Element %d failed: %s", idx, exc, exc_info=True)
                skipped += 1

        logger.info(
            "[BUILD] Inserted %d element(s), skipped %d", inserted, skipped
        )

    # ------------------------------------------------------------------
    def _build_para(self, pd: ParagraphData, doc: Document) -> Optional[Any]:
        """
        Build a <w:p> element for this paragraph.
        If an LLM-formatted text is available in llm_para_map, it is used instead of the
        raw source XML (preserving only the blueprint style and footnote references).
        Otherwise falls back to the deep-copy + rPr-clean strategy.
        """
        llm_text = self.llm_para_map.get(id(pd))
        if llm_text:
            logger.debug(
                "[BUILD] LLM para | '%s' → '%.55s'",
                pd.original_style_name, llm_text,
            )
            return self._build_para_from_llm_text(pd, doc, llm_text)

        # ── Original deep-copy path ────────────────────────────────────
        if pd.raw_xml is None:
            logger.debug("[BUILD] Para has no raw_xml – skipping")
            return None

        p_elem = copy.deepcopy(pd.raw_xml)
        _strip_tracking_attrs(p_elem)

        target_name = self.mapper.map(pd.original_style_name, pd.semantic_class, pd.heading_level)
        target_id = self._style_id(target_name, doc)

        if pd.is_section_break:
            logger.warning(
                "[BUILD] Source para has inline sectPr ('%s') – stripping it "
                "(blueprint page layout preserved)",
                pd.original_style_name,
            )

        self._reset_pPr(p_elem, target_id)
        self._clean_runs(p_elem)

        logger.debug(
            "[BUILD] Built para: '%s' → '%s' (id='%s') | fn_refs=%s",
            pd.original_style_name, target_name, target_id, pd.has_footnote_refs,
        )
        return p_elem

    # ------------------------------------------------------------------
    def _build_para_from_llm_text(
        self, pd: "ParagraphData", doc: Document, md_text: str
    ) -> Any:
        """
        Build a brand-new <w:p> element from LLM-formatted markdown text.
        The blueprint style is applied via <w:pPr>. Footnote-reference runs
        from the original source XML are re-attached at the end.
        """
        target_name = self.mapper.map(pd.original_style_name, pd.semantic_class, pd.heading_level)
        target_id   = self._style_id(target_name, doc)

        p_elem = OxmlElement("w:p")

        # Paragraph properties: only the style reference
        pPr    = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), target_id)
        pPr.append(pStyle)
        p_elem.append(pPr)

        # Runs from parsed markdown
        md_runs = parse_md_runs(md_text)
        for rd in md_runs:
            if not rd.text:
                continue
            r_elem = OxmlElement("w:r")
            if rd.bold or rd.italic:
                rPr = OxmlElement("w:rPr")
                if rd.bold:
                    rPr.append(OxmlElement("w:b"))
                    rPr.append(OxmlElement("w:bCs"))
                if rd.italic:
                    rPr.append(OxmlElement("w:i"))
                    rPr.append(OxmlElement("w:iCs"))
                r_elem.append(rPr)
            t_elem = OxmlElement("w:t")
            t_elem.text = rd.text
            if rd.text and (rd.text[0] == " " or rd.text[-1] == " "):
                t_elem.set(_XML_SPACE, "preserve")
            r_elem.append(t_elem)
            p_elem.append(r_elem)

        # Re-attach any footnote-reference runs from the original XML
        if pd.has_footnote_refs and pd.raw_xml is not None:
            for r_ref in _xpath(pd.raw_xml, ".//w:r[.//w:footnoteReference]"):
                p_elem.append(copy.deepcopy(r_ref))
            logger.debug("[BUILD] Footnote refs re-attached to LLM-built para")

        return p_elem

    # ------------------------------------------------------------------
    def _reset_pPr(self, p_elem: Any, style_id: str) -> None:
        """
        Reset paragraph properties to contain only the style reference.
        All direct formatting (indents, spacing, alignment, section breaks)
        from the source is discarded – blueprint style defines everything.
        """
        pPr = p_elem.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p_elem.insert(0, pPr)

        # Strip all existing children
        stripped = [child.tag for child in pPr]
        for child in list(pPr):
            pPr.remove(child)
        if stripped:
            logger.debug("[BUILD]   pPr stripped: %s", stripped)

        # Re-add only the style reference
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), style_id)
        pPr.append(pStyle)

    # ------------------------------------------------------------------
    def _clean_runs(self, p_elem: Any) -> None:
        """
        For each <w:r> in the paragraph:
          - If it contains a footnote reference → leave entirely untouched
          - Otherwise → strip font/color from <w:rPr>, keep KEEP_RPR_TAGS
        """
        for r_elem in p_elem.findall(qn("w:r")):
            # Preserve footnote reference runs verbatim
            fn_check = _xpath(r_elem, ".//w:footnoteReference | .//w:footnoteRef")
            if fn_check:
                logger.debug("[BUILD]   Footnote ref run preserved")
                continue

            rPr = r_elem.find(qn("w:rPr"))
            if rPr is not None:
                to_remove = [c for c in rPr if c.tag not in KEEP_RPR_TAGS]
                for child in to_remove:
                    rPr.remove(child)
                    logger.debug("[BUILD]   rPr stripped: %s", child.tag)

    # ------------------------------------------------------------------
    def _remap_table_styles(self, tbl_xml: Any, doc: Document) -> None:
        """
        For each <w:p> inside the table, remap its paragraph style to the
        blueprint equivalent.
        """
        count = 0
        for p_elem in _xpath(tbl_xml, ".//w:p"):
            try:
                # Get the source style ID from pPr/pStyle
                pStyle_elems = _xpath(p_elem, "./w:pPr/w:pStyle")
                src_id = (
                    pStyle_elems[0].get(_w("val"), "Normal")
                    if pStyle_elems
                    else "Normal"
                )
                # Resolve source style name from our ID map
                src_name = self.src_style_id_to_name.get(src_id, src_id)
                sem, hl = classify_style(src_name)
                target_name = self.mapper.map(src_name, sem, hl)
                target_id = self._style_id(target_name, doc)

                self._reset_pPr(p_elem, target_id)
                self._clean_runs(p_elem)
                count += 1
            except Exception as exc:
                logger.debug("[BUILD] Table para remap error: %s", exc)

        logger.debug("[BUILD] Table remapped: %d paragraph(s)", count)

    # ------------------------------------------------------------------
    def _style_id(self, style_name: str, doc: Document) -> str:
        """
        Return the style ID for a style name, searching:
          1. Our schema (from blueprint analysis)
          2. The live document's styles
          3. Fallback: return 'Normal'
        """
        if style_name in self.schema.styles:
            return self.schema.styles[style_name].style_id

        try:
            return doc.styles[style_name].style_id
        except Exception:
            pass

        name_lo = style_name.lower()
        for s in doc.styles:
            try:
                if s.name.lower() == name_lo:
                    return s.style_id
            except Exception:
                pass

        logger.warning(
            "[BUILD] Style '%s' not found in document – using 'Normal'", style_name
        )
        return "Normal"

    # ------------------------------------------------------------------
    def _transplant_footnotes(
        self, doc: Document, footnotes: List[FootnoteData]
    ) -> None:
        """
        Replace the blueprint copy's footnote content with the source's
        footnotes, applying the blueprint's footnote text style.
        """
        logger.info("[BUILD] Transplanting %d footnote(s)…", len(footnotes))
        try:
            fn_part = None
            for rel in doc.part.rels.values():
                if "relationships/footnotes" in rel.reltype:
                    fn_part = rel.target_part
                    break

            if fn_part is None:
                logger.warning(
                    "[BUILD] Blueprint copy has no footnotes part "
                    "– footnotes cannot be transplanted"
                )
                return

            fn_root = parse_xml(fn_part.blob)

            # Resolve blueprint footnote text style
            bp_fn_style_id = self._find_footnote_style_id(doc)
            logger.debug("[BUILD] Blueprint footnote text style id: '%s'", bp_fn_style_id)

            # Remove all existing numbered footnotes (keep id <= 0: Word internals)
            for fn_elem in _xpath(fn_root, "//w:footnote"):
                fn_id = fn_elem.get(_w("id"), "0")
                if int(fn_id) > 0:
                    fn_root.remove(fn_elem)
                    logger.debug("[BUILD] Removed blueprint footnote id=%s", fn_id)

            # Insert source footnotes with blueprint styling
            for fd in footnotes:
                fn_xml = copy.deepcopy(fd.raw_xml)
                _strip_tracking_attrs(fn_xml)

                p_elems = _xpath(fn_xml, ".//w:p")
                for para_idx, p_elem in enumerate(p_elems):
                    # ── Determine blueprint style ──────────────────────
                    pStyle_elems = _xpath(p_elem, "./w:pPr/w:pStyle")
                    src_id = (
                        pStyle_elems[0].get(_w("val"), "FootnoteText")
                        if pStyle_elems
                        else "FootnoteText"
                    )
                    src_name = self.src_style_id_to_name.get(src_id, src_id)
                    sem, _ = classify_style(src_name)

                    if sem == "footnote":
                        target_id = bp_fn_style_id
                    else:
                        target_name = self.mapper.map(src_name, sem, 0)
                        target_id = self._style_id(target_name, doc)

                    self._reset_pPr(p_elem, target_id)

                    # ── Check for LLM-formatted replacement ───────────
                    llm_text = None
                    if para_idx < len(fd.paragraphs):
                        llm_text = self.llm_fn_map.get(id(fd.paragraphs[para_idx]))

                    if llm_text:
                        # Replace all non-marker runs with LLM-built runs
                        marker_runs = _xpath(p_elem, ".//w:r[.//w:footnoteRef]")
                        for r in list(p_elem.findall(qn("w:r"))):
                            if r not in marker_runs:
                                p_elem.remove(r)
                        
                        # Apply blueprint marker formatting to the preserved marker runs
                        for r_marker in marker_runs:
                            self._apply_fn_ref_style(r_marker)

                        for rd in parse_md_runs(llm_text):
                            if not rd.text:
                                continue
                            r_elem = OxmlElement("w:r")
                            if rd.bold or rd.italic:
                                rPr = OxmlElement("w:rPr")
                                if rd.bold:
                                    rPr.append(OxmlElement("w:b"))
                                if rd.italic:
                                    rPr.append(OxmlElement("w:i"))
                                r_elem.append(rPr)
                            t_elem = OxmlElement("w:t")
                            t_elem.text = rd.text
                            if rd.text and (rd.text[0] == " " or rd.text[-1] == " "):
                                t_elem.set(_XML_SPACE, "preserve")
                            r_elem.append(t_elem)
                            p_elem.append(r_elem)
                        logger.debug(
                            "[BUILD] LLM footnote id=%s para %d: '%.50s'",
                            fd.footnote_id, para_idx, llm_text,
                        )
                    else:
                        # ── Original run-clean path ────────────────────────
                        # Apply blueprint style to <w:footnoteRef> marker run;
                        # strip source aesthetics from all other runs.
                        for r_elem in p_elem.findall(qn("w:r")):
                            fn_ref_check = _xpath(r_elem, ".//w:footnoteRef")
                            if fn_ref_check:
                                self._apply_fn_ref_style(r_elem)
                                continue
                            rPr = r_elem.find(qn("w:rPr"))
                            if rPr is not None:
                                for child in [c for c in rPr if c.tag not in KEEP_RPR_TAGS]:
                                    rPr.remove(child)

                    # Ensure separator after marker matches blueprint convention
                    # (Re-applied even for LLM text to ensure tab preservation)
                    self._normalize_fn_separator(p_elem)

                fn_root.append(fn_xml)
                logger.debug(
                    "[BUILD] Inserted footnote id=%s (%d para(s))",
                    fd.footnote_id,
                    len(fd.paragraphs),
                )

            # Commit updated XML
            updated = etree.tostring(fn_root, encoding="utf-8", xml_declaration=True)
            fn_part._blob = updated
            logger.info(
                "[BUILD] ✓ Footnote XML committed (%d footnote(s))", len(footnotes)
            )

        except Exception as exc:
            logger.error("[BUILD] Footnote transplant failed: %s", exc, exc_info=True)

    def _find_footnote_style_id(self, doc: Document) -> str:
        """Find the footnote text paragraph style ID in the document."""
        for s in doc.styles:
            sem, _ = classify_style(s.name)
            if sem == "footnote":
                return s.style_id
        return "FootnoteText"

    # ------------------------------------------------------------------
    def _apply_fn_ref_style(self, r_elem: Any) -> None:
        """
        Replace the footnoteRef marker run's rPr with the blueprint's actual
        footnote-number formatting, read directly from the blueprint's own footnotes.

        Priority:
          1. Verbatim <w:rPr> deep-copied from the blueprint's real marker runs
             (captures font name, size, vertAlign/superscript, color exactly).
          2. Fallback: a bare <w:rStyle> referencing the blueprint's detected
             FootnoteReference character style — used when the blueprint had no
             numbered footnotes to sample from.
        """
        # Remove any existing rPr first
        old_rPr = r_elem.find(qn("w:rPr"))
        if old_rPr is not None:
            r_elem.remove(old_rPr)

        if self.schema.footnote_marker_rPr_xml is not None:
            # Use the exact rPr read from the blueprint's footnotes
            new_rPr = copy.deepcopy(self.schema.footnote_marker_rPr_xml)
            r_elem.insert(0, new_rPr)
            logger.debug("[BUILD] FootnoteRef run: applied blueprint marker rPr (verbatim)")
        else:
            # Fallback: only apply the character style reference
            new_rPr = OxmlElement("w:rPr")
            rStyle = OxmlElement("w:rStyle")
            rStyle.set(qn("w:val"), self.schema.footnote_ref_char_style_id)
            new_rPr.append(rStyle)
            r_elem.insert(0, new_rPr)
            logger.debug(
                "[BUILD] FootnoteRef run: applied char style '%s' (fallback)",
                self.schema.footnote_ref_char_style_id,
            )

    # ------------------------------------------------------------------
    def _normalize_fn_separator(self, p_elem: Any) -> None:
        """
        Ensure the run immediately after <w:footnoteRef> carries the same
        separator text or tab element as the blueprint's footnotes.

        Three cases handled:
          • Separator run exists, content matches → no-op
          • Separator run exists, content differs → replace its content
          • No run after marker, blueprint wants one → insert a new run
        Only acts when schema.footnote_separator was successfully read from the blueprint.
        """
        wanted = self.schema.footnote_separator
        if wanted is None:
            return  # blueprint had no footnotes; cannot determine convention

        _XML_SPACE_ATTR = "{http://www.w3.org/XML/1998/namespace}space"
        runs = list(p_elem.findall(qn("w:r")))

        def _make_sep_run(text: str):
            sep_r = OxmlElement("w:r")
            if text == "\t":
                sep_r.append(OxmlElement("w:tab"))
            else:
                t_elem = OxmlElement("w:t")
                t_elem.text = text
                if " " in text:
                    t_elem.set(_XML_SPACE_ATTR, "preserve")
                sep_r.append(t_elem)
            return sep_r

        for ri, r_elem in enumerate(runs):
            if not _xpath(r_elem, ".//w:footnoteRef"):
                continue

            if ri + 1 < len(runs):
                next_r = runs[ri + 1]
                has_tab = next_r.find(qn("w:tab")) is not None
                t_elems = next_r.findall(qn("w:t"))
                current_text = "".join(t.text or "" for t in t_elems)
                
                # A run is a separator run if it has a tab OR is purely whitespace text
                is_sep_run = has_tab or current_text.strip() == ""

                if is_sep_run:
                    # Decide if current content matches 'wanted'
                    # (Note: we treat any existing tab element as equivalent to wanted="\t")
                    matches = (has_tab and wanted == "\t") or (not has_tab and current_text == wanted)

                    if wanted == "":
                        # Blueprint has no separator — clear the run's content
                        for child in list(next_r):
                            if child.tag in (qn("w:t"), qn("w:tab")):
                                next_r.remove(child)
                        logger.debug("[BUILD] Footnote separator cleared")
                    elif not matches:
                        # Replace all existing content with the blueprint's separator
                        for child in list(next_r):
                            if child.tag in (qn("w:t"), qn("w:tab")):
                                next_r.remove(child)
                        
                        if wanted == "\t":
                            next_r.append(OxmlElement("w:tab"))
                        else:
                            t_elem = OxmlElement("w:t")
                            t_elem.text = wanted
                            if " " in wanted:
                                t_elem.set(_XML_SPACE_ATTR, "preserve")
                            next_r.append(t_elem)
                        logger.debug(
                            "[BUILD] Footnote separator: %r → %r", 
                            ("<w:tab/>" if has_tab else current_text), 
                            wanted
                        )
                    # else: matches — no-op
                else:
                    # Next run is actual footnote text, not a separator run.
                    if wanted:
                        # Blueprint uses a separator — insert a new run before the text
                        next_r.addprevious(_make_sep_run(wanted))
                        logger.debug(
                            "[BUILD] Footnote separator inserted before text: %r", wanted
                        )
                    # else: blueprint has no separator either — nothing to do
            elif wanted:
                # No run at all after the marker — insert a new separator run
                r_elem.addnext(_make_sep_run(wanted))
                logger.debug("[BUILD] Footnote separator run appended: %r", wanted)
            break  # found the footnoteRef; done


# ============================================================================
# LLM – MULTI-PROVIDER CLIENT
# ============================================================================

# Separator used in LLM batch responses – must be on its own line
_BATCH_SEP = "---PARAGRAPH---"

# xml:space attribute for preserving spaces in <w:t>
_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


class MultiProviderLLMClient:
    """
    Unified synchronous LLM client.

    OpenAI-compatible providers (OpenAI, Nebius, Scaleway, OpenRouter, Mistral, Groq, Ollama)
    all use `openai.OpenAI(base_url=…)`.
    Anthropic uses its own SDK.
    Poe uses fastapi-poe (async, wrapped synchronously).
    """

    def complete(self, system: str, user: str, config: LLMConfig) -> str:
        """Send a chat completion and return the assistant's text."""
        # Candidate models list: primary model followed by fallbacks
        models_to_try = [config.model] + config.fallback_models
        
        last_exception = None
        
        for model_id in models_to_try:
            current_config = copy.copy(config)
            current_config.model = model_id
            
            logger.info("[LLM] %s: Trying model '%s'...", config.provider.value, model_id)
            
            for attempt in range(1, config.max_retries + 1):
                try:
                    if config.provider == LLMProvider.ANTHROPIC:
                        return self._anthropic(system, user, current_config)
                    elif config.provider == LLMProvider.POE:
                        return self._poe(system, user, current_config)
                    elif config.provider == LLMProvider.OLLAMA:
                        return self._ollama(system, user, current_config)
                    else:
                        return self._openai_compat(system, user, current_config)
                except Exception as exc:
                    last_exception = exc
                    exc_str = str(exc).lower()
                    is_rate_limit = "429" in exc_str or "rate limit" in exc_str
                    is_model_not_found = "404" in exc_str or "not found" in exc_str or "does not exist" in exc_str
                    
                    if is_model_not_found:
                        logger.warning("[LLM] %s: Model '%s' not found. Trying next fallback...", 
                                       config.provider.value, model_id)
                        break # Exit attempt loop, try next model
                    
                    # Exponential backoff: retry_delay * (2 ^ (attempt-1))
                    delay = config.retry_delay_s * (2 ** (attempt - 1))
                    header_delay = None
                    
                    # OpenAI / Groq / OpenRouter often put it in headers
                    if hasattr(exc, "response") and hasattr(exc.response, "headers"):
                        retry_after = exc.response.headers.get("retry-after")
                        if retry_after and retry_after.isdigit():
                            header_delay = float(retry_after)
                    
                    if header_delay:
                        delay = max(delay, header_delay + 1.0) # Add 1s buffer
                    elif is_rate_limit:
                        delay *= 2 # Extra patience for rate limits
                    
                    if is_rate_limit:
                        logger.warning(
                            "[LLM] %s rate limited (429) for model '%s'. Waiting %.1f seconds... (Attempt %d/%d)",
                            config.provider.value, model_id, delay, attempt, config.max_retries
                        )
                    else:
                        logger.warning(
                            "[LLM] %s model '%s' attempt %d/%d failed: %s",
                            config.provider.value, model_id, attempt, config.max_retries, exc,
                        )
                    
                    if attempt < config.max_retries:
                        time.sleep(delay)
                    else:
                        logger.error("[LLM] %s: All retries failed for model '%s'.", 
                                     config.provider.value, model_id)
            
        raise RuntimeError(
            f"[LLM] All models and retries failed for {config.provider.value}. Last error: {last_exception}"
        )

    def get_available_models(self, config: LLMConfig) -> List[Dict[str, Any]]:
        """
        Query available models from the provider's /models endpoint.
        Returns a list of model info dictionaries with parsed capabilities.
        """
        logger.info("[LLM] Querying available models for %s...", config.provider.value)
        try:
            if config.provider == LLMProvider.ANTHROPIC:
                return self._list_anthropic_models(config)
            elif config.provider == LLMProvider.POE:
                return [{"id": "Poe Bots", "capabilities": "Unknown"}]
            elif config.provider == LLMProvider.OLLAMA:
                return self._list_ollama_models(config)
            else:
                return self._list_openai_compat_models(config)
        except Exception as e:
            logger.error("[LLM] Failed to query models for %s: %s", config.provider.value, e)
            return []

    def _list_openai_compat_models(self, config: LLMConfig) -> List[Dict[str, Any]]:
        base_url = config.base_url or PROVIDER_DEFAULTS.get(config.provider.value, {}).get("base_url")
        if not base_url:
            return []
        
        headers = {"Authorization": f"Bearer {config.api_key}"}
        if config.provider == LLMProvider.OPENROUTER:
            headers["X-Title"] = "CrispTranslator"
            
        try:
            resp = requests.get(f"{base_url}/models", headers=headers, timeout=10)
            if resp.status_code != 200:
                logger.error("[LLM] HTTP %d: %s", resp.status_code, resp.text)
                return []
            
            data = resp.json()
            models = []
            raw_models = data.get("data", []) if isinstance(data, dict) else data
            
            for m in raw_models:
                m_id = m.get("id")
                if not m_id: continue
                
                # Parse capabilities
                caps = []
                if "context_window" in m:
                    caps.append(f"ctx: {m['context_window']}")
                elif "context_length" in m:
                    caps.append(f"ctx: {m['context_length']}")
                
                if m.get("pricing"):
                    p = m["pricing"]
                    caps.append(f"price: {p.get('prompt', '?')}/{p.get('completion', '?')}")
                
                info = {
                    "id": m_id,
                    "capabilities": ", ".join(caps) if caps else "Available",
                    "raw": m
                }
                models.append(info)
                logger.debug("[LLM] Found model: %s (%s)", m_id, info["capabilities"])
                
            return sorted(models, key=lambda x: x["id"])
        except Exception as e:
            logger.debug("[LLM] Model listing failed: %s", e)
            return []

    def _list_anthropic_models(self, config: LLMConfig) -> List[Dict[str, Any]]:
        # Anthropic recently added /v1/models
        headers = {
            "x-api-key": config.api_key,
            "anthropic-version": "2023-06-01"
        }
        try:
            resp = requests.get("https://api.anthropic.com/v1/models", headers=headers, timeout=10)
            if resp.status_code == 200:
                data = resp.json()
                models = []
                for m in data.get("data", []):
                    m_id = m.get("id")
                    info = {
                        "id": m_id,
                        "capabilities": f"Display: {m.get('display_name', '')}",
                        "raw": m
                    }
                    models.append(info)
                    logger.debug("[LLM] Found Anthropic model: %s", m_id)
                return models
        except:
            pass
        # Fallback if endpoint is not available
        return [{"id": "claude-3-5-sonnet-20241022", "capabilities": "Hardcoded Fallback"}]

    def _list_ollama_models(self, config: LLMConfig) -> List[Dict[str, Any]]:
        base_url = config.base_url or "http://localhost:11434/api"
        try:
            resp = requests.get(f"{base_url}/tags", timeout=5)
            if resp.status_code == 200:
                data = resp.json()
                models = []
                for m in data.get("models", []):
                    m_id = m.get("name")
                    details = m.get("details", {})
                    caps = f"{details.get('parameter_size', '?')} params, {details.get('format', '?')}"
                    models.append({"id": m_id, "capabilities": caps, "raw": m})
                    logger.debug("[LLM] Found Ollama model: %s (%s)", m_id, caps)
                return models
        except:
            pass
        return []

    # ── OpenAI-compatible ─────────────────────────────────────────────
    def _openai_compat(self, system: str, user: str, config: LLMConfig) -> str:
        if not HAS_OPENAI:
            raise ImportError("openai package not installed")
        from openai import OpenAI
        kwargs: Dict[str, Any] = {"api_key": config.api_key}
        base = config.base_url or PROVIDER_DEFAULTS.get(config.provider.value, {}).get("base_url")
        if base:
            kwargs["base_url"] = base
        # OpenRouter requires attribution headers
        extra_headers = {}
        if config.provider == LLMProvider.OPENROUTER:
            extra_headers = {
                "HTTP-Referer": "https://github.com/crisptranslator",
                "X-Title": "CrispTranslator",
            }
        client = OpenAI(**kwargs)
        logger.debug("[LLM] %s → %s | sys=%d chars user=%d chars",
                     config.provider.value, config.model, len(system), len(user))
        resp = client.chat.completions.create(
            model=config.model,
            messages=[
                {"role": "system", "content": system},
                {"role": "user",   "content": user},
            ],
            max_tokens=config.max_tokens,
            temperature=config.temperature,
            extra_headers=extra_headers or None,
        )
        text = resp.choices[0].message.content or ""
        logger.debug("[LLM] Response: %d chars", len(text))
        return text

    # ── Ollama ────────────────────────────────────────────────────────
    def _ollama(self, system: str, user: str, config: LLMConfig) -> str:
        base_url = config.base_url or "http://localhost:11434/api"
        logger.debug("[LLM] ollama → %s | sys=%d chars user=%d chars",
                     config.model, len(system), len(user))
        
        prompt = f"{system}\n\n{user}" if system else user
        
        resp = requests.post(
            f"{base_url}/generate",
            json={
                "model": config.model,
                "prompt": prompt,
                "stream": False,
                "options": {
                    "temperature": config.temperature,
                }
            },
            timeout=180
        )
        if resp.status_code != 200:
            raise RuntimeError(f"Ollama error {resp.status_code}: {resp.text}")
        
        text = resp.json().get("response", "")
        logger.debug("[LLM] Response: %d chars", len(text))
        return text

    # ── Anthropic ─────────────────────────────────────────────────────
    def _anthropic(self, system: str, user: str, config: LLMConfig) -> str:
        if not HAS_ANTHROPIC:
            raise ImportError("anthropic package not installed")
        import anthropic as ant
        client = ant.Anthropic(api_key=config.api_key)
        logger.debug("[LLM] anthropic → %s | sys=%d chars user=%d chars",
                     config.model, len(system), len(user))
        resp = client.messages.create(
            model=config.model,
            system=system,
            messages=[{"role": "user", "content": user}],
            max_tokens=config.max_tokens,
        )
        text = resp.content[0].text if resp.content else ""
        logger.debug("[LLM] Response: %d chars", len(text))
        return text

    # ── Poe ───────────────────────────────────────────────────────────
    def _poe(self, system: str, user: str, config: LLMConfig) -> str:
        if not HAS_POE:
            raise ImportError("fastapi-poe package not installed: pip install fastapi-poe")
        import fastapi_poe as fp

        combined = f"{system}\n\n{user}" if system else user

        async def _query() -> str:
            msg = fp.ProtocolMessage(role="user", content=combined)
            parts: List[str] = []
            async for partial in fp.get_bot_response(
                messages=[msg],
                bot_name=config.model,
                api_key=config.api_key,
            ):
                if isinstance(partial, fp.PartialResponse):
                    parts.append(partial.text)
            return "".join(parts)

        try:
            loop = asyncio.get_running_loop()
            # If a loop is already running, run in a new thread
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor(max_workers=1) as ex:
                fut = ex.submit(asyncio.run, _query())
                return fut.result(timeout=120)
        except RuntimeError:
            return asyncio.run(_query())


# ============================================================================
# LLM – BLUEPRINT TEXT EXTRACTION
# ============================================================================

def extract_blueprint_text(doc: Document, max_chars: int = 40_000) -> str:
    """
    Extract a representative flat-text sample from the blueprint.
    Headings are prefixed with '#'/'##'/etc. for context.
    Up to 20 footnotes are appended at the end.
    Stops when max_chars is reached.
    """
    lines: List[str] = []
    total = 0

    for para in doc.paragraphs:
        if total >= max_chars:
            break
        text = para.text.strip()
        if not text:
            lines.append("")
            continue
        _, level = classify_style(para.style.name if para.style else "Normal")
        prefix = ("#" * level + " ") if level > 0 else ""
        line = f"{prefix}{text}"
        lines.append(line)
        total += len(line)

    # Append a sample of footnotes
    try:
        fn_part = None
        for rel in doc.part.rels.values():
            if "relationships/footnotes" in rel.reltype:
                fn_part = rel.target_part
                break
        if fn_part and total < max_chars:
            root = parse_xml(fn_part.blob)
            count = 0
            for fn_elem in _xpath(root, "//w:footnote"):
                fn_id = fn_elem.get(_w("id"), "0")
                if int(fn_id) <= 0:
                    continue
                parts = []
                for t in _xpath(fn_elem, ".//w:t"):
                    parts.append(t.text or "")
                fn_text = "".join(parts).strip()
                if fn_text:
                    line = f"^[{fn_id}] {fn_text}"
                    lines.append(line)
                    total += len(line)
                    count += 1
                if count >= 20 or total >= max_chars:
                    break
    except Exception as exc:
        logger.debug("[LLM] Footnote extraction for blueprint text failed: %s", exc)

    result = "\n".join(lines)
    logger.info("[LLM] Blueprint text sample: %d chars (%d lines)", len(result), len(lines))
    return result


# ============================================================================
# LLM – STYLE GUIDE GENERATOR
# ============================================================================

_SG_SYSTEM = """\
You are an expert scholarly editor and citation specialist deriving a comprehensive editorial style guide from a document.
Your output must be a precise instruction set for reformatting text to match this document's exact standards.

USER PRIORITY RULE:
If the user provides supplementary style information, those rules take ABSOLUTE PRECEDENCE over patterns you observe in the excerpt.

MANDATORY AREAS OF ANALYSIS:
1. CITATION STYLE: Meticulously analyze footnote citations. Identify patterns for:
   - Book/article titles (italic? quotes?)
   - Author names (Full name? Surname? All caps? Roman?)
   - Volume/Issue/Page notation (S. 12? p. 12? 12-15? 12f?)
   - Punctuation between components (Commas? Colons? Slashes?)
   - Repeated citations (Vgl.? See? Ibid.? ebenda?)
2. PUNCTUATION & SYMBOLS: Identify specific choices for:
   - Quotation marks (»...«, „...“, "...", '...')
   - Dashes (— em-dash, – en-dash)
   - Spaces before/after symbols
3. NAMES & TERMS: Identify treatment of personal names, institutional names, and foreign terms.

Write the style guide as actionable, imperative rules (e.g., "Always use...", "Never italicize...").
"""

_SG_USER_TMPL = """\
Below is a comprehensive excerpt from the **blueprint document**, including sampled footnotes. 
Analyse its editorial conventions with extreme care.

DOCUMENT EXCERPT (Body & Footnotes):
──────────────────────────────────────────────────
{blueprint_text}
──────────────────────────────────────────────────
{extra_section}

Produce a **MASTER STYLE GUIDE** in Markdown. 

CRITICAL: Your guide must be detailed enough to handle complex academic citations and specific punctuation (like »...« quotation marks) without ambiguity.

Structure your guide:
1. **Absolute User Overrides** (Include any rules from the 'Additional information' section here first)
2. **Language & Register**
3. **Personal & Institutional Names**
4. **Foreign-Language Terms & Transliteration**
5. **Inline Emphasis & Special Symbols** (Meticulously specify quotation marks: » vs „ vs ")
6. **Footnote & Citation System** (Provide specific templates for books, articles, and repeats)
7. **Heading & Layout Conventions**

Aim for scholarly perfection.
"""


class StyleGuideGenerator:
    """
    Phase 1-LLM: Extracts a text sample from the blueprint and asks the LLM
    to produce a self-contained editorial style guide (Markdown).
    """

    def __init__(self, client: MultiProviderLLMClient) -> None:
        self.client = client

    def generate(
        self,
        blueprint_doc: Document,
        extra_paths: Optional[List[Path]],
        config: LLMConfig,
    ) -> str:
        logger.info("[LLM-SG] Generating style guide from blueprint…")
        blueprint_text = extract_blueprint_text(blueprint_doc, config.blueprint_context_chars)

        extra_section = ""
        if extra_paths:
            parts = []
            for p in extra_paths:
                try:
                    content = p.read_text(encoding="utf-8", errors="replace")
                    parts.append(f"### Additional style information from '{p.name}':\n{content}")
                    logger.info("[LLM-SG] Loaded extra styleguide: %s (%d chars)", p.name, len(content))
                except Exception as exc:
                    logger.warning("[LLM-SG] Could not read extra styleguide '%s': %s", p, exc)
            if parts:
                extra_section = (
                    "\nIn addition, here is supplementary style information provided by the user:\n\n"
                    + "\n\n".join(parts)
                    + "\n\n"
                )

        user_msg = _SG_USER_TMPL.format(
            blueprint_text=blueprint_text,
            extra_section=extra_section,
        )
        styleguide = self.client.complete(_SG_SYSTEM, user_msg, config)
        logger.info("[LLM-SG] Style guide generated: %d chars", len(styleguide))
        logger.debug("[LLM-SG] Style guide preview:\n%s", styleguide[:600])
        return styleguide


# ============================================================================
# LLM – MARKDOWN RUN PARSER
# ============================================================================

# Matches inline Markdown in priority order (longest markers first)
_MD_TOKEN = re.compile(
    r"\*\*\*(.+?)\*\*\*"   # bold+italic  → group 1
    r"|\*\*(.+?)\*\*"       # bold         → group 2
    r"|\*(.+?)\*"           # italic (*)   → group 3
    r"|_(.+?)_"             # italic (_)   → group 4
    r"|([^*_\n]+)"          # plain text   → group 5
    r"|([*_]+|\n)",         # stray chars  → group 6
    re.DOTALL,
)


def parse_md_runs(text: str) -> List["RunData"]:
    """
    Convert a string with Markdown inline formatting into a list of RunData.
    Handles ***bold+italic***, **bold**, *italic*, _italic_, plain text.
    Stray asterisks/underscores are emitted as plain runs.
    """
    runs: List[RunData] = []
    for m in _MD_TOKEN.finditer(text):
        g1, g2, g3, g4, g5, g6 = m.groups()
        if g1:
            runs.append(RunData(text=g1, bold=True,  italic=True))
        elif g2:
            runs.append(RunData(text=g2, bold=True))
        elif g3:
            runs.append(RunData(text=g3, italic=True))
        elif g4:
            runs.append(RunData(text=g4, italic=True))
        elif g5:
            runs.append(RunData(text=g5))
        elif g6:
            runs.append(RunData(text=g6))   # stray marker as plain text
    return [r for r in runs if r.text]


# ============================================================================
# LLM – CONTENT FORMATTER
# ============================================================================

_FMT_SYSTEM = """\
You are a scholarly editor applying a strict editorial style guide to existing text.
Your task is to re-format the provided text to match the Style Guide's exact conventions.

CONSTRAINTS:
1. SUBSTANTIVE VERBATIM: Do NOT change the substantive meaning, names, or titles. 
2. EDITORIAL RE-FORMATTING: You MUST change punctuation, quotation marks, and citation structure (e.g., brackets vs commas, colons vs spaces) to strictly follow the Style Guide.
3. DO NOT translate, summarize, or paraphrase.
4. DO NOT add any introductory remarks or commentary.

Use Markdown for inline formatting:
  *italic*          for italic text
  **bold**          for bold text
  ***bold italic*** for bold + italic
No other Markdown. Return only the re-formatted paragraph text.
Return EXACTLY one response for each input paragraph.
"""

_PARA_USER_TMPL = """\
STYLE GUIDE:
──────────────────────────────────────────────────
{styleguide}
──────────────────────────────────────────────────

Apply this style guide to each of the {n} paragraphs below.
Return EXACTLY {n} formatted paragraphs separated by the line:
{sep}
Do NOT number them. Do NOT add any commentary or blank lines between the separator and the next paragraph.

PARAGRAPHS:
{content}
"""

_FN_USER_TMPL = """\
STYLE GUIDE:
──────────────────────────────────────────────────
{styleguide}
──────────────────────────────────────────────────

Apply this style guide to each of the {n} footnotes below.
Footnotes often contain citations, names, foreign terms and references —
pay special attention to the citation and name conventions in the style guide.
Return EXACTLY {n} formatted footnotes separated by the line:
{sep}
Do NOT number them. Do NOT add commentary.

FOOTNOTES:
{content}
"""


class LLMContentFormatter:
    """
    Phase 2-LLM: Sends batches of paragraphs / footnotes to the LLM with
    the generated style guide, and parses the response back to plain strings
    (with Markdown inline markers).
    """

    def __init__(self, client: MultiProviderLLMClient) -> None:
        self.client = client

    # ------------------------------------------------------------------
    def format_paragraphs(
        self,
        paras: List["ParagraphData"],
        styleguide: str,
        config: LLMConfig,
    ) -> Dict[int, str]:
        """
        Format a list of paragraphs. Returns {id(pd): formatted_text}.
        Falls back to the original text on LLM failure.
        """
        return self._format_batch(paras, styleguide, config, mode="para")

    def format_footnotes(
        self,
        footnotes: List["FootnoteData"],
        styleguide: str,
        config: LLMConfig,
    ) -> Dict[int, str]:
        """
        Format all footnote paragraphs. Returns {id(para_data): formatted_text}.
        """
        # Flatten footnote paragraphs
        flat: List["ParagraphData"] = []
        for fd in footnotes:
            flat.extend(fd.paragraphs)
        return self._format_batch(flat, styleguide, config, mode="footnote")

    # ------------------------------------------------------------------
    def _format_batch(
        self,
        paras: List["ParagraphData"],
        styleguide: str,
        config: LLMConfig,
        mode: str,
    ) -> Dict[int, str]:
        result: Dict[int, str] = {}
        # Only format paragraphs that have actual text
        to_format = [p for p in paras if p.get_text().strip()]
        logger.info(
            "[LLM-FMT] Formatting %d %s(s) in batches of %d…",
            len(to_format), mode, config.para_batch_size,
        )

        for batch_start in range(0, len(to_format), config.para_batch_size):
            # Inter-batch delay to stay under rate limits
            if batch_start > 0:
                batch_delay = 2.0 # 2 seconds between batches
                if config.provider == LLMProvider.GROQ:
                    batch_delay = 15.0 # Extra delay for Groq (very tight limits)
                logger.info("[LLM-FMT] Inter-batch delay: %.1fs...", batch_delay)
                time.sleep(batch_delay)

            batch = to_format[batch_start: batch_start + config.para_batch_size]
            texts = [p.get_text() for p in batch]

            content = f"\n{_BATCH_SEP}\n".join(texts)
            tmpl = _FN_USER_TMPL if mode == "footnote" else _PARA_USER_TMPL
            user_msg = tmpl.format(
                styleguide=styleguide,
                n=len(batch),
                sep=_BATCH_SEP,
                content=content,
            )

            logger.debug(
                "[LLM-FMT] Batch %d–%d (%d items), user_msg=%d chars",
                batch_start, batch_start + len(batch) - 1, len(batch), len(user_msg),
            )

            try:
                response = self.client.complete(_FMT_SYSTEM, user_msg, config)
                parsed = self._parse_response(response, len(batch), texts)
            except Exception as exc:
                logger.error("[LLM-FMT] Batch failed, using originals: %s", exc)
                parsed = texts

            for pd, formatted in zip(batch, parsed):
                if formatted.strip():
                    result[id(pd)] = formatted
                    logger.debug(
                        "[LLM-FMT] Para formatted: orig='%.50s' → fmt='%.50s'",
                        pd.get_text(), formatted,
                    )

        return result

    # ------------------------------------------------------------------
    @staticmethod
    def _parse_response(response: str, expected: int, originals: List[str]) -> List[str]:
        """
        Split the LLM response on _BATCH_SEP and return exactly `expected` strings.
        Falls back to originals for any missing entries.
        """
        parts = [p.strip() for p in response.split(_BATCH_SEP)]
        parts = [p for p in parts if p]   # remove empties

        if len(parts) != expected:
            logger.warning(
                "[LLM-FMT] Expected %d parts, got %d — padding/truncating",
                expected, len(parts),
            )
        # Pad with originals if too short, truncate if too long
        while len(parts) < expected:
            parts.append(originals[len(parts)])
        return parts[:expected]


# ============================================================================
# MAIN ORCHESTRATOR
# ============================================================================


class FormatTransplanter:
    """
    Orchestrates the four-phase format transplant pipeline:
      Phase 1: Analyse blueprint → BlueprintSchema
      Phase 2: Extract source content → ParagraphData / FootnoteData
      Phase 3: Build style map → StyleMapper
      Phase 4: Assemble output document → DocumentBuilder
    """

    def run(
        self,
        blueprint_path: Path,
        source_path: Path,
        output_path: Path,
        user_style_overrides: Optional[Dict[str, str]] = None,
    ) -> None:
        logger.info("═" * 60)
        logger.info("FORMAT TRANSPLANT")
        logger.info("  Blueprint : %s", blueprint_path)
        logger.info("  Source    : %s", source_path)
        logger.info("  Output    : %s", output_path)
        if user_style_overrides:
            logger.info("  Overrides : %s", user_style_overrides)
        logger.info("═" * 60)

        # Phase 1 ─────────────────────────────────────────────────────
        logger.info("Phase 1 – Analysing blueprint…")
        bp_doc = Document(str(blueprint_path))
        analyzer = BlueprintAnalyzer()
        schema = analyzer.analyze(bp_doc)

        # Phase 2 ─────────────────────────────────────────────────────
        logger.info("Phase 2 – Extracting source content…")
        src_doc = Document(str(source_path))
        extractor = ContentExtractor()
        body_elements, footnotes = extractor.extract(src_doc)

        # Phase 3 ─────────────────────────────────────────────────────
        logger.info("Phase 3 – Building style map…")
        mapper = StyleMapper(schema, user_style_overrides)
        mapper.log_full_table(body_elements)

        # Phase 4 ─────────────────────────────────────────────────────
        logger.info("Phase 4 – Building output document…")
        builder = DocumentBuilder(schema, mapper)
        builder.src_style_id_to_name = extractor.src_style_id_to_name
        builder.build(blueprint_path, output_path, body_elements, footnotes)

        logger.info("═" * 60)
        logger.info("✓ Format transplant complete → %s", output_path)
        logger.info("═" * 60)


# ============================================================================
# LLM FORMAT TRANSPLANTER
# ============================================================================


class LLMFormatTransplanter:
    """
    Extends the base transplant pipeline with two LLM phases:

      Phase 1-LLM  StyleGuideGenerator  — blueprint excerpt → styleguide.md
      Phase 2-LLM  LLMContentFormatter  — source paragraphs/footnotes → formatted markdown

    The four structural phases (blueprint analysis, content extraction, style
    mapping, document assembly) are unchanged; DocumentBuilder picks up the
    LLM-formatted text automatically via its llm_para_map / llm_fn_map.

    llm_mode controls which parts go through the LLM:
      "both"           — paragraphs and footnotes (default)
      "paragraphs"     — body paragraphs only
      "footnotes"      — footnotes only
      "styleguide_only"— generate (and optionally save) styleguide, no output doc
    """

    def run(
        self,
        blueprint_path: Path,
        source_path: Path,
        output_path: Path,
        llm_config: LLMConfig,
        extra_styleguide_paths: Optional[List[Path]] = None,
        styleguide_in: Optional[Path] = None,
        styleguide_out: Optional[Path] = None,
        llm_mode: str = "both",
        user_style_overrides: Optional[Dict[str, str]] = None,
        debug_limit: Optional[int] = None,
    ) -> Optional[Path]:
        """
        Returns the path to the saved styleguide if styleguide_out was set,
        otherwise None.
        """
        logger.info("═" * 60)
        logger.info("LLM FORMAT TRANSPLANT")
        logger.info("  Blueprint  : %s", blueprint_path)
        logger.info("  Source     : %s", source_path)
        logger.info("  Output     : %s", output_path)
        logger.info("  Provider   : %s / %s", llm_config.provider.value, llm_config.model)
        logger.info("  LLM mode   : %s", llm_mode)
        if debug_limit:
            logger.info("  Debug limit: %d paragraphs", debug_limit)
        logger.info("  Batch size : %d  Context chars: %d",
                    llm_config.para_batch_size, llm_config.blueprint_context_chars)
        logger.info("═" * 60)

        client = MultiProviderLLMClient()

        # ── Phase 1: Blueprint analysis ────────────────────────────────
        logger.info("Phase 1 – Analysing blueprint…")
        bp_doc = Document(str(blueprint_path))
        schema = BlueprintAnalyzer().analyze(bp_doc)

        # ── Phase 1-LLM: Styleguide generation / loading ───────────────
        if styleguide_in and styleguide_in.exists():
            styleguide_md = styleguide_in.read_text(encoding="utf-8")
            logger.info("Phase 1-LLM – Loaded existing styleguide from %s (%d chars)",
                        styleguide_in, len(styleguide_md))
        else:
            logger.info("Phase 1-LLM – Generating style guide…")
            sg_gen = StyleGuideGenerator(client)
            styleguide_md = sg_gen.generate(bp_doc, extra_styleguide_paths, llm_config)

        saved_sg: Optional[Path] = None
        if styleguide_out:
            styleguide_out.write_text(styleguide_md, encoding="utf-8")
            saved_sg = styleguide_out
            logger.info("Phase 1-LLM – Style guide saved → %s", styleguide_out)

        if llm_mode == "styleguide_only":
            logger.info("Mode = styleguide_only — stopping after style guide generation.")
            return saved_sg

        # ── Phase 2: Content extraction ────────────────────────────────
        logger.info("Phase 2 – Extracting source content…")
        src_doc = Document(str(source_path))
        extractor = ContentExtractor()
        body_elements, footnotes = extractor.extract(src_doc)

        # Apply debug limit if requested
        if debug_limit:
            count = 0
            limited_body = []
            for e in body_elements:
                limited_body.append(e)
                if e.semantic_class != "table":
                    count += 1
                if count >= debug_limit:
                    break
            body_elements = limited_body
            logger.info("Debug limit applied: only processing first %d body paragraphs.", debug_limit)

        # ── Phase 2-LLM: LLM content formatting ───────────────────────
        formatter = LLMContentFormatter(client)
        llm_para_map: Dict[int, str] = {}
        llm_fn_map:   Dict[int, str] = {}

        if llm_mode in ("both", "paragraphs"):
            logger.info("Phase 2-LLM – Formatting body paragraphs…")
            para_candidates = [
                e for e in body_elements
                if e.semantic_class != "table" and e.get_text().strip()
            ]
            llm_para_map = formatter.format_paragraphs(para_candidates, styleguide_md, llm_config)
            logger.info("Phase 2-LLM – %d paragraph(s) formatted by LLM", len(llm_para_map))

        if llm_mode in ("both", "footnotes") and footnotes:
            logger.info("Phase 2-LLM – Formatting footnotes…")
            llm_fn_map = formatter.format_footnotes(footnotes, styleguide_md, llm_config)
            logger.info("Phase 2-LLM – %d footnote paragraph(s) formatted by LLM", len(llm_fn_map))

        # ── Phase 3: Style mapping ─────────────────────────────────────
        logger.info("Phase 3 – Building style map…")
        mapper = StyleMapper(schema, user_style_overrides)
        mapper.log_full_table(body_elements)

        # ── Phase 4: Document assembly ─────────────────────────────────
        logger.info("Phase 4 – Building output document…")
        builder = DocumentBuilder(schema, mapper)
        builder.src_style_id_to_name = extractor.src_style_id_to_name
        builder.llm_para_map = llm_para_map
        builder.llm_fn_map   = llm_fn_map
        builder.build(blueprint_path, output_path, body_elements, footnotes)

        logger.info("═" * 60)
        logger.info("✓ LLM format transplant complete → %s", output_path)
        logger.info("═" * 60)
        return saved_sg


# ============================================================================
# CLI
# ============================================================================


def _parse_overrides(pairs: Optional[List[str]]) -> Dict[str, str]:
    """Parse 'Source Style=Blueprint Style' strings from the CLI."""
    if not pairs:
        return {}
    result: Dict[str, str] = {}
    for item in pairs:
        if "=" not in item:
            logger.warning("Ignoring invalid --style-map entry (no '='): '%s'", item)
            continue
        src, _, bp = item.partition("=")
        result[src.strip()] = bp.strip()
    return result


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Format Transplant – apply blueprint formatting to source document content",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Basic structural transplant (no LLM)
  python format_transplant.py blueprint.docx source.docx output.docx

  # LLM style pass with Nebius Llama
  python format_transplant.py blueprint.docx source.docx output.docx \\
      --llm nebius --llm-model meta-llama/Meta-Llama-3.1-70B-Instruct

  # Generate and inspect the style guide first, then re-use it
  python format_transplant.py blueprint.docx source.docx output.docx \\
      --llm anthropic --llm-mode styleguide_only --styleguide-out style.md
  # (edit style.md if needed, then:)
  python format_transplant.py blueprint.docx source.docx output.docx \\
      --llm anthropic --styleguide-in style.md

  # Extra styleguide documents + OpenRouter
  python format_transplant.py blueprint.docx source.docx output.docx \\
      --llm openrouter --llm-model anthropic/claude-opus-4-5 \\
      --extra-styleguide house_rules.md dmg_table.txt

Debug tips:
  python format_transplant.py ... -v 2>&1 | tee run.log
  grep "\\[MAPPER\\]"  run.log   # style mapping
  grep "\\[LLM\\]"     run.log   # LLM calls
  grep "\\[LLM-FMT\\]" run.log   # batch formatting
        """,
    )
    # ── Positional ─────────────────────────────────────────────────────
    parser.add_argument("blueprint", help="Blueprint DOCX – provides all formatting")
    parser.add_argument("source",    help="Source DOCX – provides all text content")
    parser.add_argument("output",    help="Output DOCX path")

    # ── General ────────────────────────────────────────────────────────
    parser.add_argument("-v", "--verbose", action="store_true",
                        help="Enable DEBUG logging")
    parser.add_argument("--style-map", nargs="+", metavar="SRC=BP",
                        help='Style overrides: "Source Style=Blueprint Style"')

    # ── LLM options ────────────────────────────────────────────────────
    llm_group = parser.add_argument_group("LLM options (all optional)")
    llm_group.add_argument(
        "--llm",
        choices=list(PROVIDER_DEFAULTS.keys()),
        default=None,
        metavar="PROVIDER",
        help="LLM provider: " + ", ".join(PROVIDER_DEFAULTS.keys()),
    )
    llm_group.add_argument("--llm-model",   default=None, metavar="MODEL",
                           help="Model name (default: provider default)")
    llm_group.add_argument("--llm-key",     default=None, metavar="KEY",
                           help="API key (default: read from env var)")
    llm_group.add_argument(
        "--llm-mode",
        choices=["both", "paragraphs", "footnotes", "styleguide_only"],
        default="both",
        help="Which content goes through LLM (default: both)",
    )
    llm_group.add_argument("--styleguide-out", default=None, metavar="PATH",
                           help="Save generated style guide to this .md file")
    llm_group.add_argument("--styleguide-in",  default=None, metavar="PATH",
                           help="Load pre-existing style guide (skip generation)")
    llm_group.add_argument("--extra-styleguide", nargs="+", default=None, metavar="PATH",
                           help="Extra style-info files sent to LLM during generation")
    llm_group.add_argument("--llm-context-chars", type=int, default=40_000, metavar="N",
                           help="Blueprint chars to send for styleguide gen (default: 40000)")
    llm_group.add_argument("--llm-batch", type=int, default=15, metavar="N",
                           help="Source paragraphs per LLM batch (default: 15)")
    llm_group.add_argument("--debug-limit", type=int, default=None, metavar="N",
                           help="Process only first N paragraphs (for testing)")

    args = parser.parse_args()

    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
        logger.debug("DEBUG logging enabled")

    blueprint_path = Path(args.blueprint)
    source_path    = Path(args.source)
    output_path    = Path(args.output)

    if not blueprint_path.exists():
        logger.error("Blueprint file not found: %s", blueprint_path)
        sys.exit(1)
    if not source_path.exists():
        logger.error("Source file not found: %s", source_path)
        sys.exit(1)
    if output_path.exists():
        logger.warning("Output file already exists – will overwrite: %s", output_path)

    overrides = _parse_overrides(args.style_map)

    # ── LLM path ──────────────────────────────────────────────────────
    if args.llm:
        try:
            cfg = llm_config_from_args(args.llm, args.llm_model, args.llm_key)
            cfg.blueprint_context_chars = args.llm_context_chars
            cfg.para_batch_size         = args.llm_batch
        except ValueError as exc:
            logger.error("%s", exc)
            sys.exit(1)

        extra_sg = [Path(p) for p in args.extra_styleguide] if args.extra_styleguide else None
        sg_in    = Path(args.styleguide_in)  if args.styleguide_in  else None
        sg_out   = Path(args.styleguide_out) if args.styleguide_out else None

        transplanter = LLMFormatTransplanter()
        try:
            transplanter.run(
                blueprint_path=blueprint_path,
                source_path=source_path,
                output_path=output_path,
                llm_config=cfg,
                extra_styleguide_paths=extra_sg,
                styleguide_in=sg_in,
                styleguide_out=sg_out,
                llm_mode=args.llm_mode,
                user_style_overrides=overrides,
                debug_limit=args.debug_limit,
            )
        except Exception as exc:
            logger.error("Fatal error: %s", exc, exc_info=True)
            sys.exit(1)

    # ── Structural-only path ──────────────────────────────────────────
    else:
        transplanter = FormatTransplanter()
        try:
            transplanter.run(blueprint_path, source_path, output_path, overrides)
        except Exception as exc:
            logger.error("Fatal error: %s", exc, exc_info=True)
            sys.exit(1)


if __name__ == "__main__":
    main()
